# -*- coding: utf-8 -*-
"""Dựng file báo giá gửi khách (Word .docx và PDF) bằng thuần Python.

Nút "Xuất / Gửi báo giá" trên form Báo giá chạy vào đây. Word dựng bằng
python-docx, PDF bằng reportlab — KHÔNG cần wkhtmltopdf.

🔴 Cả hai file chỉ in GIÁ BÁN cho khách. Tuyệt đối không in giá vốn (vật tư/
công đoạn/giá sàn/markup) — đó là số nội bộ.

Tách khỏi dl_quotation.py cho gọn, và import thư viện đặt TRONG hàm để máy nào
thiếu thư viện thì chỉ lỗi lúc bấm xuất file, không chặn nạp cả module."""

import base64
import io

from odoo import models, _
from odoo.exceptions import UserError
from odoo.tools.misc import formatLang, format_date, file_path

# Nhóm font đăng ký cho reportlab (dùng lại nếu đã đăng ký trong tiến trình).
_PDF_FONT = "DLSans"
_PDF_FONT_BOLD = "DLSans-Bold"


class DlQuotationDocument(models.Model):
    _inherit = "dl.quotation"

    # ------------------------------------------------------------------
    # Gom sẵn dữ liệu để Word và PDF in ra giống hệt nhau (một nguồn duy nhất).
    # ------------------------------------------------------------------
    def _document_context(self):
        """Trả về dict dữ liệu đã format sẵn (dòng hàng, tiền, điều khoản) cho
        cả hai hàm dựng Word/PDF dùng chung."""
        self.ensure_one()
        company = self.company_id or self.env.company
        partner = self.partner_id
        currency = self.currency_id or company.currency_id

        def money(value):
            return formatLang(self.env, value or 0.0, currency_obj=currency)

        lines = []
        for idx, line in enumerate(self.line_ids, start=1):
            lines.append({
                "stt": idx,
                "name": line.name or "",
                "qty": line.qty,
                "qty_txt": formatLang(self.env, line.qty or 0.0),
                "price_unit": money(line.price_unit),
                "subtotal": money(line.price_subtotal),
            })

        try:
            amount_words = currency.amount_to_text(self.amount_total)
        except Exception:  # noqa: BLE001 — không chặn xuất file nếu đọc chữ lỗi
            amount_words = ""

        # Điều khoản gửi khách — chỉ đưa vào file những mục có nhập.
        terms = [(label, text) for label, text in (
            (_("Thanh toán"), self.payment_terms),
            (_("Giao hàng"), self.delivery_terms),
            (_("Bảo hành"), self.warranty_terms),
        ) if text]

        return {
            "company": company,
            "partner": partner,
            "currency": currency,
            "lines": lines,
            "money": money,
            "amount_words": amount_words,
            "terms": terms,
        }

    # ==================================================================
    # WORD (.docx) — python-docx
    # ==================================================================
    def _build_docx(self):
        """Dựng file Word báo giá, trả về bytes."""
        self.ensure_one()
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise UserError(_(
                "Thiếu thư viện 'python-docx' để xuất Word. Cài bằng: "
                "pip install python-docx"))

        ctx = self._document_context()
        company, partner = ctx["company"], ctx["partner"]
        doc = Document()

        # Font mặc định.
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        # --- Đầu trang: logo + công ty ---
        # Logo chèn ở khổ 4,5 cm ngang; thiếu logo thì bỏ qua, không chặn xuất file.
        if company.logo:
            logo_p = doc.add_paragraph()
            logo_p.paragraph_format.space_after = Pt(2)
            logo_p.add_run().add_picture(
                io.BytesIO(base64.b64decode(company.logo)), width=Cm(4.5))

        p = doc.add_paragraph()
        run = p.add_run(company.name or "")
        run.bold = True
        run.font.size = Pt(14)
        info = []
        addr = ", ".join(filter(None, [company.street, company.city]))
        if addr:
            info.append("Địa chỉ: " + addr)
        if company.vat:
            info.append("Mã số thuế: " + company.vat)
        contact = " | ".join(filter(None, [
            ("ĐT: " + company.phone) if company.phone else "",
            ("Email: " + company.email) if company.email else "",
        ]))
        for text in filter(None, info + ([contact] if contact else [])):
            cp = doc.add_paragraph(text)
            cp.paragraph_format.space_after = Pt(0)

        # --- Tiêu đề ---
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        trun = title.add_run("BÁO GIÁ")
        trun.bold = True
        trun.font.size = Pt(20)
        title.paragraph_format.space_before = Pt(10)
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srun = sub.add_run("Số: %s%s" % (
            self.name or "",
            (" — Phiên bản %s" % self.revision) if self.revision > 1 else ""))
        srun.bold = True

        # --- Khách hàng + thông tin báo giá ---
        meta = doc.add_table(rows=1, cols=2)
        meta.autofit = True
        left, right = meta.rows[0].cells
        lp = left.paragraphs[0]
        lp.add_run("Kính gửi: ").bold = True
        lp.add_run(partner.name or "")
        for text in filter(None, [
            ("Mã số thuế: " + partner.vat) if partner.vat else "",
            ("Địa chỉ: " + ", ".join(filter(None, [partner.street, partner.city])))
            if partner.street else "",
            ("Điện thoại: " + partner.phone) if partner.phone else "",
            ("Email: " + partner.email) if partner.email else "",
        ]):
            left.add_paragraph(text)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.add_run("Ngày báo giá: %s" % (
            format_date(self.env, self.date_order) if self.date_order else ""))
        if self.validity_date:
            rv = right.add_paragraph("Hiệu lực đến: %s" % format_date(
                self.env, self.validity_date))
            rv.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph(
            "Công ty chúng tôi trân trọng gửi tới Quý khách bảng báo giá cho "
            "các hạng mục sau:")

        # --- Bảng dòng ---
        headers = ["Số thứ tự", "Nội dung / Hạng mục", "Số lượng",
                   "Đơn giá", "Thành tiền"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, text in enumerate(headers):
            run = hdr[i].paragraphs[0].add_run(text)
            run.bold = True
        for ln in ctx["lines"]:
            cells = table.add_row().cells
            cells[0].text = str(ln["stt"])
            cells[1].text = ln["name"]
            cells[2].text = ln["qty_txt"]
            cells[3].text = ln["price_unit"]
            cells[4].text = ln["subtotal"]
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # --- Tổng hợp ---
        money = ctx["money"]
        totals = [("Cộng tiền hàng:", money(self.amount_untaxed), False)]
        if self.discount_pct:
            totals.append(("Chiết khấu (%g%%):" % self.discount_pct,
                           "-" + money(self.discount_amount), False))
            totals.append(("Sau chiết khấu:", money(self.amount_before_vat), False))
        if self.vat_pct:
            totals.append(("Thuế GTGT (%g%%):" % self.vat_pct,
                           money(self.vat_amount), False))
        totals.append(("TỔNG THANH TOÁN:", money(self.amount_total), True))
        tt = doc.add_table(rows=0, cols=2)
        tt.alignment = WD_TABLE_ALIGNMENT.RIGHT
        for label, value, strong in totals:
            row = tt.add_row().cells
            lr = row[0].paragraphs[0].add_run(label)
            vp = row[1].paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            vr = vp.add_run(value)
            lr.bold = vr.bold = strong

        if ctx["amount_words"]:
            wp = doc.add_paragraph()
            wr = wp.add_run("Bằng chữ: %s" % ctx["amount_words"])
            wr.italic = True

        # --- Điều khoản gửi khách ---
        if ctx["terms"]:
            tp = doc.add_paragraph()
            tp.add_run("ĐIỀU KHOẢN:").bold = True
            tp.paragraph_format.space_after = Pt(0)
            for label, text in ctx["terms"]:
                lp = doc.add_paragraph()
                lp.add_run("• %s: " % label).bold = True
                lp.add_run(text)
                lp.paragraph_format.space_after = Pt(0)

        if self.note:
            np = doc.add_paragraph()
            np.add_run("Ghi chú:").bold = True
            doc.add_paragraph(self.note)

        # --- Chữ ký ---
        doc.add_paragraph()
        sign = doc.add_table(rows=1, cols=2)
        sc = sign.rows[0].cells
        for cell, label in ((sc[0], "KHÁCH HÀNG"), (sc[1], "ĐẠI DIỆN CÔNG TY")):
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(label).bold = True
            note = cell.add_paragraph("(Ký, ghi rõ họ tên)")
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note.runs[0].italic = True

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ==================================================================
    # PDF — reportlab (platypus)
    # ==================================================================
    def _register_pdf_font(self):
        """Đăng ký font hỗ trợ tiếng Việt cho reportlab (một lần/tiến trình).

        Ưu tiên Montserrat đi kèm Odoo (có trên mọi bản cài, Windows & Linux),
        rồi tới font hệ thống. reportlab mặc định (Helvetica) KHÔNG render được
        dấu tiếng Việt nên bắt buộc nhúng TTF."""
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        if _PDF_FONT in pdfmetrics.getRegisteredFontNames():
            return
        import os
        reg = bold = None
        # 1) Montserrat đi kèm module web của Odoo.
        try:
            reg = file_path(
                "web/static/fonts/google/Montserrat/Montserrat-Regular.ttf")
            bold = file_path(
                "web/static/fonts/google/Montserrat/Montserrat-Bold.ttf")
        except (FileNotFoundError, ValueError):
            reg = bold = None
        # 2) Font hệ thống (Windows Arial / Linux DejaVu).
        if not reg:
            for r, b in (
                (r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf"),
                ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                 "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            ):
                if os.path.exists(r) and os.path.exists(b):
                    reg, bold = r, b
                    break
        if not reg:
            raise UserError(_(
                "Không tìm thấy font hỗ trợ tiếng Việt để tạo PDF. Vui lòng "
                "dùng định dạng Word (.docx)."))
        pdfmetrics.registerFont(TTFont(_PDF_FONT, reg))
        pdfmetrics.registerFont(TTFont(_PDF_FONT_BOLD, bold))
        pdfmetrics.registerFontFamily(
            _PDF_FONT, normal=_PDF_FONT, bold=_PDF_FONT_BOLD)

    def _build_pdf(self):
        """Dựng file PDF báo giá, trả về bytes."""
        self.ensure_one()
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import mm
            from reportlab.lib import colors
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_RIGHT
            from reportlab.platypus import (
                SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer)
        except ImportError:
            raise UserError(_(
                "Thiếu thư viện 'reportlab' để xuất PDF. Cài bằng: "
                "pip install reportlab"))

        self._register_pdf_font()
        ctx = self._document_context()
        company, partner = ctx["company"], ctx["partner"]
        money = ctx["money"]

        base = ParagraphStyle(
            "dl_base", fontName=_PDF_FONT, fontSize=10, leading=14)
        bold = ParagraphStyle("dl_bold", parent=base, fontName=_PDF_FONT_BOLD)
        right = ParagraphStyle("dl_right", parent=base, alignment=TA_RIGHT)
        right_b = ParagraphStyle(
            "dl_right_b", parent=bold, alignment=TA_RIGHT)
        center = ParagraphStyle("dl_center", parent=base, alignment=TA_CENTER)
        # `leading` phải khai lại theo `fontSize`: kế thừa từ `base` là 14pt,
        # chữ 20pt sẽ đè lên dòng ngay dưới.
        title = ParagraphStyle(
            "dl_title", parent=bold, fontSize=20, leading=25,
            alignment=TA_CENTER, spaceBefore=6, spaceAfter=2)

        story = []
        # --- Đầu trang công ty (logo + thông tin, dựng ở dl_base để 3 loại
        #     chứng từ gửi ra ngoài dùng chung một khuôn) ---
        story += company._dlm_pdf_letterhead(
            name_style=ParagraphStyle(
                "dl_co", parent=bold, fontSize=13, leading=17),
            body_style=base, width=174 * mm)

        story.append(Paragraph("BÁO GIÁ", title))
        story.append(Paragraph("Số: %s%s" % (
            self.name or "",
            (" — Phiên bản %s" % self.revision) if self.revision > 1 else ""),
            center))
        story.append(Spacer(1, 6 * mm))

        # --- Khách hàng + meta (2 cột) ---
        cust = [Paragraph("<b>Kính gửi:</b> %s" % (partner.name or ""), base)]
        for text in filter(None, [
            ("Mã số thuế: " + partner.vat) if partner.vat else "",
            ("Địa chỉ: " + ", ".join(filter(None, [partner.street, partner.city])))
            if partner.street else "",
            ("Điện thoại: " + partner.phone) if partner.phone else "",
            ("Email: " + partner.email) if partner.email else "",
        ]):
            cust.append(Paragraph(text, base))
        meta = [Paragraph("Ngày báo giá: %s" % (
            format_date(self.env, self.date_order) if self.date_order else ""),
            right)]
        if self.validity_date:
            meta.append(Paragraph("Hiệu lực đến: %s" % format_date(
                self.env, self.validity_date), right))
        head_tbl = Table([[cust, meta]], colWidths=[110 * mm, 60 * mm])
        head_tbl.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ]))
        story.append(head_tbl)
        story.append(Spacer(1, 3 * mm))
        story.append(Paragraph(
            "Công ty chúng tôi trân trọng gửi tới Quý khách bảng báo giá cho "
            "các hạng mục sau:", base))
        story.append(Spacer(1, 2 * mm))

        # --- Bảng dòng ---
        data = [[Paragraph("Số thứ tự", ParagraphStyle("h", parent=bold, alignment=TA_CENTER, textColor=colors.white)),
                 Paragraph("Nội dung / Hạng mục", ParagraphStyle("h2", parent=bold, textColor=colors.white)),
                 Paragraph("Số lượng", ParagraphStyle("h3", parent=bold, alignment=TA_CENTER, textColor=colors.white)),
                 Paragraph("Đơn giá", ParagraphStyle("h4", parent=bold, alignment=TA_RIGHT, textColor=colors.white)),
                 Paragraph("Thành tiền", ParagraphStyle("h5", parent=bold, alignment=TA_RIGHT, textColor=colors.white))]]
        for ln in ctx["lines"]:
            data.append([
                Paragraph(str(ln["stt"]), center),
                Paragraph(ln["name"], base),
                Paragraph(ln["qty_txt"], center),
                Paragraph(ln["price_unit"], right),
                Paragraph(ln["subtotal"], right),
            ])
        line_tbl = Table(
            data, colWidths=[18 * mm, 70 * mm, 21 * mm, 31 * mm, 31 * mm],
            repeatRows=1)
        line_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#34495e")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#b0b0b0")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(line_tbl)
        story.append(Spacer(1, 3 * mm))

        # --- Tổng hợp (căn phải) ---
        totals = [["Cộng tiền hàng:", money(self.amount_untaxed)]]
        if self.discount_pct:
            totals.append(["Chiết khấu (%g%%):" % self.discount_pct,
                           "-" + money(self.discount_amount)])
            totals.append(["Sau chiết khấu:", money(self.amount_before_vat)])
        if self.vat_pct:
            totals.append(["Thuế GTGT (%g%%):" % self.vat_pct,
                           money(self.vat_amount)])
        tot_rows = [[Paragraph(l, right), Paragraph(v, right)]
                    for l, v in totals]
        tot_rows.append([Paragraph("TỔNG THANH TOÁN:", right_b),
                         Paragraph(money(self.amount_total), right_b)])
        tot_tbl = Table(tot_rows, colWidths=[110 * mm, 60 * mm])
        tot_tbl.setStyle(TableStyle([
            ("LINEABOVE", (0, -1), (-1, -1), 1, colors.black),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        story.append(tot_tbl)

        if ctx["amount_words"]:
            story.append(Spacer(1, 2 * mm))
            story.append(Paragraph(
                "<i>Bằng chữ: %s</i>" % ctx["amount_words"], base))

        # --- Điều khoản gửi khách ---
        if ctx["terms"]:
            story.append(Spacer(1, 3 * mm))
            story.append(Paragraph("<b>ĐIỀU KHOẢN:</b>", base))
            for label, text in ctx["terms"]:
                story.append(Paragraph(
                    "<b>%s:</b> %s" % (label, (text or "").replace("\n", "<br/>")),
                    base))

        if self.note:
            story.append(Spacer(1, 3 * mm))
            story.append(Paragraph("<b>Ghi chú:</b>", base))
            story.append(Paragraph(
                (self.note or "").replace("\n", "<br/>"), base))

        # --- Chữ ký ---
        story.append(Spacer(1, 12 * mm))
        sign = Table([[
            [Paragraph("KHÁCH HÀNG", ParagraphStyle("sk", parent=bold, alignment=TA_CENTER)),
             Paragraph("<i>(Ký, ghi rõ họ tên)</i>", center)],
            [Paragraph("ĐẠI DIỆN CÔNG TY", ParagraphStyle("sc", parent=bold, alignment=TA_CENTER)),
             Paragraph("<i>(Ký, ghi rõ họ tên)</i>", center)],
        ]], colWidths=[85 * mm, 85 * mm])
        sign.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
        story.append(sign)

        buf = io.BytesIO()
        SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=18 * mm, rightMargin=18 * mm,
            topMargin=15 * mm, bottomMargin=15 * mm,
            title="Bao gia %s" % (self.name or ""),
        ).build(story)
        return buf.getvalue()

    # ==================================================================
    # Đóng gói: sinh nội dung + tạo ir.attachment + đăng chatter
    # ==================================================================
    def _quotation_doc_filename(self, doc_format):
        self.ensure_one()
        ext = "pdf" if doc_format == "pdf" else "docx"
        return "Bao_gia_%s.%s" % ((self.name or "moi").replace("/", "_"), ext)

    def _generate_quotation_document(self, doc_format):
        """Trả về (bytes nội dung, mimetype) theo định dạng chọn."""
        self.ensure_one()
        if doc_format == "pdf":
            return self._build_pdf(), "application/pdf"
        return (self._build_docx(),
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document")

    def _create_quotation_attachment(self, doc_format):
        """Sinh file báo giá và lưu thành ir.attachment gắn vào báo giá."""
        self.ensure_one()
        content, mimetype = self._generate_quotation_document(doc_format)
        return self.env["ir.attachment"].create({
            "name": self._quotation_doc_filename(doc_format),
            "datas": base64.b64encode(content),
            "res_model": "dl.quotation",
            "res_id": self.id,
            "mimetype": mimetype,
        })

    def _post_quotation_document_to_chatter(self, doc_format="pdf"):
        """Kẹp file báo giá vào chatter làm bằng chứng "đã gửi khách". Gọi tự
        động lúc bấm Gửi khách. Bọc try/except vì đây chỉ là bằng chứng phụ —
        thiếu thư viện/font làm hỏng file cũng KHÔNG được chặn việc gửi."""
        self.ensure_one()
        try:
            content, _mime = self._generate_quotation_document(doc_format)
        except Exception:  # noqa: BLE001 — mọi lỗi dựng file đều không chặn gửi
            return
        self.message_post(
            body=_("Đã phát hành báo giá — đính kèm bản gửi khách."),
            attachments=[(self._quotation_doc_filename(doc_format), content)],
        )

    # ------------------------------------------------------------------
    # Nút "Xuất / Gửi báo giá" trên form → mở dialog chọn Word/PDF.
    # ------------------------------------------------------------------
    def action_open_export_wizard(self):
        """Mở dialog chọn định dạng file rồi tải về / gửi email."""
        self.ensure_one()
        if not self.line_ids:
            raise UserError(_("Báo giá chưa có dòng nào để xuất."))
        # 🔴 Cổng chặn gửi khách phải bịt CẢ đường này. Xuất được file rồi gửi
        # tay cho khách là đi vòng qua cổng — cổng còn lại chỉ là hình thức.
        if self.dlm_send_blocked:
            raise UserError(_(
                "Chưa xuất được file gửi khách: giá phần vật tư phải mua thêm "
                "chưa được Mua hàng xác nhận. Con số trên báo giá lúc này là "
                "TẠM TÍNH theo bảng giá nhà cung cấp; xuất ra rồi gửi tay cho "
                "khách thì vẫn là cam kết trên một con số chưa ai kiểm. Bấm "
                "“Hỏi giá nhà cung cấp” trước."))
        return {
            "type": "ir.actions.act_window",
            "name": _("Xuất / Gửi báo giá"),
            "res_model": "dl.quotation.export.wizard",
            "view_mode": "form",
            "target": "new",
            "context": {"default_quotation_id": self.id},
        }
