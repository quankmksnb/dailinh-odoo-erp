from docx import Document
import re
path = r"D:\FPTU\do_van_an\Documents\Reports\Report 4.0_TDS.docx"
doc = Document(path)
out = open(r"D:\FPTU\do_van_an\dailinh-odoo-erp\.tmp_tds\cell.txt", "w", encoding="utf-8")

t = doc.tables[3]
out.write("=== ROW 1 (first body row) XML ===\n")
out.write(t.rows[1]._tr.xml[:3500])

out.write("\n\n=== A 'normal' PARAGRAPH XML (from body) ===\n")
for p in doc.paragraphs:
    if p.text.strip().startswith("Cơ chế kế thừa Odoo"):
        out.write(p._p.xml[:2500])
        break

out.write("\n\n=== A 'Heading 3' PARAGRAPH XML ===\n")
for p in doc.paragraphs:
    if p.style is not None and p.style.name == "Heading 3" and p.text.strip():
        out.write(p._p.xml[:2000])
        break

out.write("\n\n=== A 'Heading 2' PARAGRAPH XML ===\n")
for p in doc.paragraphs:
    if p.style is not None and p.style.name == "Heading 2" and p.text.strip():
        out.write(p._p.xml[:2000])
        break
out.close()
print("ok")
