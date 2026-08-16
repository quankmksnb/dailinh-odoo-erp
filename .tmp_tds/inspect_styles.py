from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
import re

path = r"D:\FPTU\do_van_an\Documents\Reports\Report 4.0_TDS.docx"
doc = Document(path)
out = open(r"D:\FPTU\do_van_an\dailinh-odoo-erp\.tmp_tds\styles.txt", "w", encoding="utf-8")

out.write("=== PARAGRAPH STYLES ===\n")
for s in doc.styles:
    try:
        out.write(f"  {s.type} :: {s.name!r}  (styleId={s.style_id})\n")
    except Exception as e:
        out.write(f"  ? {e}\n")

out.write("\n=== TABLE STYLES USED IN BODY ===\n")
for t in doc.tables:
    st = t.style.name if t.style is not None else None
    hdr = " | ".join(c.text.strip()[:18] for c in t.rows[0].cells)
    out.write(f"  style={st!r:35} {len(t.rows)}x{len(t.columns)}  {hdr[:80]}\n")

out.write("\n=== SAMPLE TABLE XML (first data table, tblPr + first row) ===\n")
t = doc.tables[3]
xml = t._tbl.xml
out.write(xml[:4000])

out.write("\n\n=== SECTION PAGE WIDTH ===\n")
for s in doc.sections:
    out.write(f"  page_width={s.page_width} left={s.left_margin} right={s.right_margin} "
              f"usable={s.page_width - s.left_margin - s.right_margin}\n")
out.close()
print("ok")
