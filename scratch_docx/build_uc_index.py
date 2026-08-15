# -*- coding: utf-8 -*-
import copy
import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

SRC = 'scratch_docx/edited7b.docx'
OUT = 'scratch_docx/edited8.docx'

d = docx.Document(SRC)


def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)


REPORT_KEYWORDS = ['Xem ', 'Xem/', 'Tra cứu', 'Theo dõi', 'Giám sát', 'Tổng quan', 'Bảng tổng quan', 'Danh sách']

SYSTEM_ASSISTED = {
    'UC-010', 'UC-057', 'UC-067', 'UC-069A', 'UC-078', 'UC-079', 'UC-080',
    'UC-091', 'UC-107', 'UC-069',
}

uc_table = d.tables[2]
rows_data = []
for row in uc_table.rows[1:]:
    cells = [c.text.strip() for c in row.cells]
    uc_id, name = cells[0], cells[1]
    perms = cells[2:]
    if all(p == '—' for p in perms):
        typ = 'System'
    elif '(Planned' in name:
        typ = 'Planned'
    elif any(kw in name for kw in REPORT_KEYWORDS):
        typ = 'UI / Report'
    elif uc_id in SYSTEM_ASSISTED:
        typ = 'UI / System-assisted'
    else:
        typ = 'UI'
    rows_data.append([uc_id, name, typ])

print('Total UC rows for index:', len(rows_data))

# Find "3.4 Use Case Index" heading, then the Type legend table right after it,
# then insert the new index table right after that legend table.
body = d.element.body
heading_34 = None
for child in body.iterchildren():
    if child.tag == qn('w:p'):
        p = Paragraph(child, d)
        if p.text.strip() == '3.4 Use Case Index':
            heading_34 = p
            break
assert heading_34 is not None

# walk forward from heading to find the Type legend table (header: Type | Meaning)
node = heading_34._p
legend_table_elm = None
while True:
    node = node.getnext()
    if node is None:
        break
    if node.tag == qn('w:tbl'):
        t = Table(node, d)
        if t.rows[0].cells[0].text.strip() == 'Type':
            legend_table_elm = node
        break  # first table after heading should be the legend; stop regardless
assert legend_table_elm is not None

ref_style = Table(legend_table_elm, d).style

header = ['UC-ID', 'Use Case', 'Type']
new_table = d.add_table(rows=len(rows_data) + 1, cols=3)
new_table.style = ref_style
for c, val in enumerate(header):
    set_cell_text(new_table.rows[0].cells[c], val)
for r, values in enumerate(rows_data, start=1):
    for c, val in enumerate(values):
        set_cell_text(new_table.rows[r].cells[c], val)

tbl_elm = new_table._tbl
tbl_elm.getparent().remove(tbl_elm)
legend_table_elm.addnext(tbl_elm)

print('UC Index table inserted with', len(rows_data), 'rows.')

d.save(OUT)
print('Saved', OUT)
