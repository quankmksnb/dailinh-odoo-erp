# -*- coding: utf-8 -*-
import copy
import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

SRC = 'scratch_docx/edited6.docx'
OUT = 'scratch_docx/edited7.docx'

d = docx.Document(SRC)


def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)


def find_row_idx(table, key, col=0):
    for i, row in enumerate(table.rows):
        if row.cells[col].text.strip() == key:
            return i
    raise ValueError(key)


def clone_row_after(table, ref_row_idx, new_cell_texts):
    ref_tr = table.rows[ref_row_idx]._tr
    new_tr = copy.deepcopy(ref_tr)
    ref_tr.addnext(new_tr)
    new_row = None
    for r in table.rows:
        if r._tr is new_tr:
            new_row = r
            break
    for cell, text in zip(new_row.cells, new_cell_texts):
        set_cell_text(cell, text)
    return new_row


def find_heading(text):
    body = d.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            p = Paragraph(child, d)
            if p.text.strip() == text:
                return p
    raise ValueError('heading not found: ' + text)


# ============ 1. FT table module INV (g) ============
inv_heading = find_heading('g. Module: Kho & Mua hàng (INV)')
inv_ft_table = Table(inv_heading._p.getnext(), d)

row = inv_ft_table.rows[find_row_idx(inv_ft_table, 'FT-81')]
set_cell_text(row.cells[1], 'Chuyển kho nội bộ')
set_cell_text(row.cells[3], 'Ghi nhận chuyển vật tư/hàng hoá giữa các khu vực trong kho; có 2 nút tạo nhanh cho 2 tuyến thường dùng (Vật tư & hàng thương mại → Xưởng sản xuất, → Kho thành phẩm)')
set_cell_text(row.cells[4], 'UC-092')

row = inv_ft_table.rows[find_row_idx(inv_ft_table, 'FT-83')]
set_cell_text(row.cells[1], 'Kiểm kê và điều chỉnh tồn kho')
set_cell_text(row.cells[3], 'Đếm thực tế và điều chỉnh số lượng tồn kho cho khớp; sinh chứng từ điều chỉnh tự động khi áp dụng')
set_cell_text(row.cells[6], 'Tự phát triển (trên nền chế độ kiểm kê tồn kho gốc của Odoo)')

new_ft_rows = [
    ['FT-98', 'Hàng đợi phiếu kho', 'Should',
     'Thủ kho đăng nhập vào thẳng danh sách các phiếu đang chờ xử lý (nhận/kiểm/chuyển/giao), bấm vào từng dòng mở đúng màn xử lý tương ứng thay vì phải nhớ vào đúng menu',
     'UC-090A', '', 'Tự phát triển'],
    ['FT-99', 'Tạo và xử lý phiếu giao hàng khách', 'Must',
     'Tạo phiếu giao hàng trực tiếp từ đơn bán hàng đã xác nhận, lấy hàng từ Kho thành phẩm; đối chiếu số lượng cần giao/đã giao/còn lại; không tạo trùng phiếu khi bấm nhiều lần',
     'UC-092A', '', 'Tự phát triển'],
    ['FT-100', 'Bán phế liệu thu hồi', 'Should',
     'Cân và ghi nhận số lượng phế liệu thu hồi được từ khu vực sản xuất vào tồn kho; chọn dòng còn tồn và lập phiếu bán cho khách mua phế liệu theo giá bán đã niêm yết',
     'UC-092B', '', 'Tự phát triển'],
    ['FT-101', 'Đối chiếu dự toán và thực tế thu hồi phế liệu', 'Should',
     'So sánh khối lượng phế liệu dự toán (tính từ tỷ lệ hao hụt trên định mức của các đơn hàng đã chốt trong kỳ) với khối lượng thực tế đã cân được, theo từng tháng, để phát hiện định mức hao hụt đặt sai hoặc thất thoát',
     'UC-092C', '', 'Tự phát triển'],
]
last_row_tr = inv_ft_table.rows[-1]._tr
for row_data in new_ft_rows:
    new_tr = copy.deepcopy(last_row_tr)
    last_row_tr.addnext(new_tr)
    new_row = None
    for r in inv_ft_table.rows:
        if r._tr is new_tr:
            new_row = r
            break
    for cell, text in zip(new_row.cells, row_data):
        set_cell_text(cell, text)
    last_row_tr = new_tr

print('INV FT table rebuilt.')

# ============ 2. SCR table ============
scr_table = None
for t in d.tables:
    if t.rows[0].cells[0].text.strip() == 'SCR-ID':
        scr_table = t
        break
assert scr_table is not None

# SCR-44A: Picking Queue, inserted before SCR-45
idx_44 = find_row_idx(scr_table, 'SCR-44')
clone_row_after(scr_table, idx_44, [
    'SCR-44A', 'Picking Queue', 'INV',
    'Hàng đợi phiếu kho đang chờ Thủ kho xử lý (nhận/kiểm/chuyển/giao), bấm vào 1 dòng để mở đúng màn xử lý',
    'FT-98', 'UC-090A',
])

# SCR-46: Internal Transfer (was "Goods Issue")
idx_46 = find_row_idx(scr_table, 'SCR-46')
row_46 = scr_table.rows[idx_46]
set_cell_text(row_46.cells[1], 'Internal Transfer')
set_cell_text(row_46.cells[3], 'Ghi nhận chuyển kho nội bộ giữa các khu vực trong kho')
set_cell_text(row_46.cells[4], 'FT-81')
set_cell_text(row_46.cells[5], 'UC-092')

# SCR-46A, 46B, 46C: mới, chèn sau SCR-46
clone_row_after(scr_table, idx_46, [
    'SCR-46A', 'Customer Delivery', 'INV',
    'Tạo và xử lý phiếu giao hàng cho khách từ đơn bán hàng đã xác nhận',
    'FT-99', 'UC-092A',
])
idx_46a = find_row_idx(scr_table, 'SCR-46A')
clone_row_after(scr_table, idx_46a, [
    'SCR-46B', 'Scrap Sale', 'INV',
    'Cân, ghi nhận và bán phế liệu thu hồi từ khu vực sản xuất',
    'FT-100', 'UC-092B',
])
idx_46b = find_row_idx(scr_table, 'SCR-46B')
clone_row_after(scr_table, idx_46b, [
    'SCR-46C', 'Scrap Recovery Reconciliation', 'INV',
    'Đối chiếu khối lượng phế liệu dự toán theo định mức với khối lượng thực tế đã cân được, theo từng tháng',
    'FT-101', 'UC-092C',
])

# SCR-48: Stock Count and Adjustment — bỏ "chưa triển khai"
idx_48 = find_row_idx(scr_table, 'SCR-48')
row_48 = scr_table.rows[idx_48]
set_cell_text(row_48.cells[3], 'Đối chiếu và điều chỉnh số lượng trong kho')
set_cell_text(row_48.cells[4], 'FT-83')
set_cell_text(row_48.cells[5], 'UC-094')

print('SCR table rebuilt.')

d.save(OUT)
print('Saved', OUT)
