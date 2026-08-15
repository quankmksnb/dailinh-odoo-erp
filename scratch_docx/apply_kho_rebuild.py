# -*- coding: utf-8 -*-
import copy
import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

SRC = 'scratch_docx/edited5.docx'
OUT = 'scratch_docx/edited6.docx'

d = docx.Document(SRC)


def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)


def find_row_idx(table, key, col=0):
    for i, row in enumerate(table.rows):
        if row.cells[col].text.strip() == key:
            return i
    raise ValueError(key)


def clone_row_after(table, ref_row_idx, new_cell_texts):
    ref_tr = table.rows[ref_row_idx]._tr
    new_tr = copy.deepcopy(ref_tr)
    ref_tr.addnext(new_tr)
    new_row = None
    for r in table.rows:
        if r._tr is new_tr:
            new_row = r
            break
    for cell, text in zip(new_row.cells, new_cell_texts):
        set_cell_text(cell, text)
    return new_row


uc_table = d.tables[2]

# ---- UC-090A: Hàng đợi phiếu kho (mới) ----
idx_090 = find_row_idx(uc_table, 'UC-090')
clone_row_after(uc_table, idx_090, [
    'UC-090A', 'Hàng đợi phiếu kho (landing riêng cho Thủ kho)',
    'No', 'No', 'No', 'No', 'No', 'Full', 'No',
])

# ---- UC-092: đổi thành "Chuyển kho nội bộ" ----
idx_092 = find_row_idx(uc_table, 'UC-092')
row_092 = uc_table.rows[idx_092]
set_cell_text(row_092.cells[1], 'Chuyển kho nội bộ')
set_cell_text(row_092.cells[7], 'Full')  # Internal Accountant (đã Full sẵn, giữ nguyên)

# ---- UC-092A, UC-092B, UC-092C: mới, chèn sau UC-092 ----
clone_row_after(uc_table, idx_092, [
    'UC-092A', 'Tạo và xử lý phiếu giao hàng khách',
    'No', 'Restricted', 'Restricted', 'Restricted', 'No', 'Full', 'No',
])
idx_092a = find_row_idx(uc_table, 'UC-092A')
clone_row_after(uc_table, idx_092a, [
    'UC-092B', 'Bán phế liệu thu hồi',
    'No', 'No', 'No', 'No', 'No', 'Full', 'No',
])
idx_092b = find_row_idx(uc_table, 'UC-092B')
clone_row_after(uc_table, idx_092b, [
    'UC-092C', 'Đối chiếu dự toán và thực tế thu hồi phế liệu',
    'No', 'Full', 'No', 'No', 'Restricted', 'Full', 'No',
])

# ---- UC-093: sửa Technician No -> Restricted, gọn lại tên ----
idx_093 = find_row_idx(uc_table, 'UC-093')
row_093 = uc_table.rows[idx_093]
set_cell_text(row_093.cells[1], 'Xem tồn kho hiện tại (cảnh báo tồn tối thiểu: xem JOB-04, chưa triển khai)')
set_cell_text(row_093.cells[6], 'Restricted')  # Technician No -> Restricted (có menu, chỉ đọc)

# ---- UC-094: Kiểm kê kho -> đã triển khai, đổi tên & mô tả ----
idx_094 = find_row_idx(uc_table, 'UC-094')
row_094 = uc_table.rows[idx_094]
set_cell_text(row_094.cells[1], 'Kiểm kê và điều chỉnh tồn kho')
set_cell_text(row_094.cells[7], 'Full')  # Internal Accountant giữ Full

# ---- UC-095, UC-096, UC-097: gọn lại thành "Planned" (backlog, chưa triển khai) ----
idx_095 = find_row_idx(uc_table, 'UC-095')
row_095 = uc_table.rows[idx_095]
set_cell_text(row_095.cells[1], 'Kiểm tra khả dụng vật tư theo BOM (Planned — backlog giai đoạn sau, ATP)')

idx_096 = find_row_idx(uc_table, 'UC-096')
row_096 = uc_table.rows[idx_096]
set_cell_text(row_096.cells[1], 'Nhập kho thành phẩm theo BOM (Planned — backlog Lệnh sản xuất/B2, đã seed sẵn loại phiếu Xuất vật tư SX/Nhập thành phẩm, chưa có màn hình)')

idx_097 = find_row_idx(uc_table, 'UC-097')
row_097 = uc_table.rows[idx_097]
set_cell_text(row_097.cells[1], 'Tính nhu cầu vật tư cho phiếu nhập kho theo BOM (Planned — gắn với UC-096)')

print('Permission matrix Kho section rebuilt. Total rows now:', len(uc_table.rows))

d.save(OUT)
print('Saved', OUT)
