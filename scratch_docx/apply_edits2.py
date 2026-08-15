# -*- coding: utf-8 -*-
import copy
import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

SRC = 'scratch_docx/edited.docx'
OUT = 'scratch_docx/edited2.docx'

d = docx.Document(SRC)


def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)


def clone_row_after(table, ref_row_idx, new_cell_texts):
    ref_tr = table.rows[ref_row_idx]._tr
    new_tr = copy.deepcopy(ref_tr)
    ref_tr.addnext(new_tr)
    new_row = None
    for r in table.rows:
        if r._tr is new_tr:
            new_row = r
            break
    for cell, text in zip(new_row.cells, new_cell_texts):
        set_cell_text(cell, text)
    return new_row


def find_row_idx(table, key):
    for i, row in enumerate(table.rows):
        if row.cells[0].text.strip() == key:
            return i
    raise ValueError(key)


# ---------- 1. Actor table: Internal Accountant description ----------
actor_table = d.tables[1]
idx_acc = find_row_idx(actor_table, 'Internal Accountant')
row_acc = actor_table.rows[idx_acc]
old_desc = row_acc.cells[1].text.strip()
new_desc = (
    'Quản lý hồ sơ nhà cung cấp và bảng giá vật tư (cập nhật, gửi duyệt, đánh dấu '
    '"đang áp dụng"); kiểm tra vật tư chưa có giá; kiêm nhiệm quản lý kho: lập phiếu '
    'nhận hàng, kiểm tra chất lượng hàng nhận (QC) và theo dõi tồn kho (phần xuất kho, '
    'chuyển kho nội bộ và kiểm kê dự kiến triển khai ở giai đoạn sau).'
)
set_cell_text(row_acc.cells[1], new_desc)
print('Actor row updated. Old:', old_desc[:60], '...')

# ---------- 2. Rebuild the INV FT table (module g) with correct, non-colliding FT-IDs ----------
body = d.element.body
heading_p = None
for child in body.iterchildren():
    if child.tag == qn('w:p'):
        p = Paragraph(child, d)
        if p.text.strip() == 'g. Module: Kho & Mua hàng (INV)':
            heading_p = p
            break
assert heading_p is not None

# the table we inserted previously is the first w:tbl sibling right after this paragraph
old_ft_table_elm = heading_p._p.getnext()
assert old_ft_table_elm is not None and old_ft_table_elm.tag == qn('w:tbl'), 'expected table right after heading'
old_ft_table = Table(old_ft_table_elm, d)
ref_style = old_ft_table.style

ft_rows = [
    ['FT-ID', 'Feature', 'Priority', 'Description', 'UC-ID', 'BF-ID', 'SOURCE'],
    ['FT-75', 'Thiết lập layout kho', 'Must',
     'Hệ thống dựng sẵn 1 kho, 3 khu vực (Nhận hàng/Xưởng/Thành phẩm) và các loại phiếu kho khi cài đặt module, không có màn hình cấu hình riêng cho người dùng',
     '—', '', 'Tự phát triển'],
    ['FT-80', 'Nhận hàng NCC (bước 1/2)', 'Must',
     'Ghi nhận hàng về từ nhà cung cấp vào khu vực chờ kiểm; tự cấp số lô cho vật tư/bán thành phẩm theo dõi theo lô nếu chưa nhập; xác nhận phiếu tự sinh phiếu Kiểm & cất hàng kế tiếp',
     'UC-091', '', 'Tự phát triển (trên nền stock của Odoo)'],
    ['FT-80A', 'Kiểm tra chất lượng hàng nhận & cất hàng (QC, bước 2/2)', 'Must',
     'Nhập số lượng đạt/loại cho từng dòng hàng khi cất kho; bắt buộc lý do khi có hàng loại; chặn xác nhận nếu số liệu không hợp lệ hoặc thiếu lô',
     'UC-091A', '', 'Tự phát triển'],
    ['FT-80B', 'Tự động tách và tạo phiếu trả hàng NCC', 'Must',
     'Khi có hàng bị loại ở bước kiểm, hệ thống tự tách phần hàng lỗi và tạo phiếu trả nhà cung cấp ở trạng thái nháp, giao cho bộ phận Mua hàng xử lý',
     'UC-091B', '', 'Tự phát triển'],
    ['FT-81', 'Ghi nhận xuất kho', 'Must',
     'Chưa triển khai — dự kiến K6-K8 (giao hàng khách, xuất huỷ, xuất sản xuất thủ công)',
     'UC-092', '', '—'],
    ['FT-82', 'Xem tồn kho hiện tại', 'Must',
     'Danh sách tồn kho theo sản phẩm/lô/vị trí, chỉ xem, không hiển thị giá vốn',
     'UC-093', '', 'Tự phát triển (trên nền tồn kho của Odoo)'],
    ['FT-83', 'Kiểm kê kho', 'Should',
     'Chưa triển khai — hiện dựa vào chế độ điều chỉnh tồn kho gốc của Odoo, không có màn hình riêng của DLM-ERP',
     'UC-094', '', '—'],
    ['FT-84', 'Truy vết lô hàng', 'Should',
     'Mỗi lô lưu lại nhà cung cấp, ngày nhận và phiếu nhập gốc; hiển thị trên màn Tồn kho và xem lại được từ màn Lô hàng riêng',
     'UC-093', '', 'Tự phát triển'],
]

new_table = d.add_table(rows=len(ft_rows), cols=7)
new_table.style = ref_style
for r, values in enumerate(ft_rows):
    for c, val in enumerate(values):
        set_cell_text(new_table.rows[r].cells[c], val)

new_tbl_elm = new_table._tbl
new_tbl_elm.getparent().remove(new_tbl_elm)
old_ft_table_elm.addprevious(new_tbl_elm)
old_ft_table_elm.getparent().remove(old_ft_table_elm)
print('FT table (module g) rebuilt with', len(ft_rows) - 1, 'rows.')

# ---------- 3. Fix SCR table UC-ID / FT-ID references (SCR-45..48) + add SCR-45A, SCR-47A ----------
scr_table = d.tables[9]

idx_45 = find_row_idx(scr_table, 'SCR-45')
row_45 = scr_table.rows[idx_45]
set_cell_text(row_45.cells[4], 'FT-80')
set_cell_text(row_45.cells[5], 'UC-091')

row_45a = clone_row_after(scr_table, idx_45, [
    'SCR-45A', 'Goods Receipt Inspection (QC)', 'INV',
    'Kiểm tra chất lượng hàng nhận về: nhập số lượng đạt/loại, lý do loại; hệ thống tự tạo phiếu trả nhà cung cấp khi có hàng lỗi',
    'FT-80A, FT-80B', 'UC-091A, UC-091B',
])

idx_46 = find_row_idx(scr_table, 'SCR-46')
row_46 = scr_table.rows[idx_46]
set_cell_text(row_46.cells[3], 'Ghi nhận phiếu xuất kho (chưa triển khai — kế hoạch K6-K8)')
set_cell_text(row_46.cells[4], 'FT-81')
set_cell_text(row_46.cells[5], 'UC-092')

idx_47 = find_row_idx(scr_table, 'SCR-47')
row_47 = scr_table.rows[idx_47]
set_cell_text(row_47.cells[4], 'FT-82, FT-84')
set_cell_text(row_47.cells[5], 'UC-093')

row_47a = clone_row_after(scr_table, idx_47, [
    'SCR-47A', 'Lot List', 'INV',
    'Danh sách lô hàng, truy vết nhà cung cấp và phiếu nhập gốc của từng lô',
    'FT-84', 'UC-093',
])

idx_48 = find_row_idx(scr_table, 'SCR-48')
row_48 = scr_table.rows[idx_48]
set_cell_text(row_48.cells[3], 'Đối chiếu và điều chỉnh số lượng trong kho (chưa triển khai — kế hoạch K6-K8)')
set_cell_text(row_48.cells[4], 'FT-83')
set_cell_text(row_48.cells[5], 'UC-094')

print('SCR table fixed.')

d.save(OUT)
print('Saved', OUT)
