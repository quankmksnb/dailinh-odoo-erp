# -*- coding: utf-8 -*-
import copy
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell

SRC = 'scratch_docx/original.docx'
OUT = 'scratch_docx/edited.docx'

d = docx.Document(SRC)

# ---------- helpers ----------

def set_cell_text(cell, text):
    # clear all paragraphs, keep first paragraph's formatting/style as a base
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    # remove any extra paragraphs beyond the first
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)


def clone_row_after(table, ref_row_idx, new_cell_texts):
    ref_tr = table.rows[ref_row_idx]._tr
    new_tr = copy.deepcopy(ref_tr)
    ref_tr.addnext(new_tr)
    # wrap into python-docx Row-like access via re-reading table
    # find the row object corresponding to new_tr
    new_row = None
    for r in table.rows:
        if r._tr is new_tr:
            new_row = r
            break
    for cell, text in zip(new_row.cells, new_cell_texts):
        set_cell_text(cell, text)
    return new_row

# ---------- 1. Permission matrix table (UC section) ----------
uc_table = d.tables[2]

# UC-091 row (index 97 in body overall, but table row index -> find by scanning)

def find_row_idx(table, uc_id):
    for i, row in enumerate(table.rows):
        if row.cells[0].text.strip() == uc_id:
            return i
    raise ValueError(uc_id)

idx_091 = find_row_idx(uc_table, 'UC-091')
row_091 = uc_table.rows[idx_091]
set_cell_text(row_091.cells[1], 'Ghi nhận nhận hàng NCC (bước 1/2)')
set_cell_text(row_091.cells[8], 'Restricted')  # Purchasing Staff: Full -> Restricted

# Insert UC-091A right after UC-091
row_091a = clone_row_after(uc_table, idx_091, [
    'UC-091A',
    'Kiểm tra chất lượng hàng nhận & cất hàng (QC)',
    'No', 'No', 'No', 'No', 'No', 'Full', 'No',
])

# Insert UC-091B right after UC-091A
idx_091a = find_row_idx(uc_table, 'UC-091A')
row_091b = clone_row_after(uc_table, idx_091a, [
    'UC-091B',
    'Xử lý phiếu trả hàng NCC (khi QC phát hiện hàng lỗi)',
    'No', 'No', 'No', 'No', 'No', 'Restricted', 'Full',
])

# UC-092: annotate as not yet implemented
idx_092 = find_row_idx(uc_table, 'UC-092')
row_092 = uc_table.rows[idx_092]
set_cell_text(row_092.cells[1], 'Ghi nhận xuất kho (chưa triển khai — kế hoạch K6-K8)')

# UC-093: drop the min-stock-alert phrase from the title, note it's tracked by JOB-04
idx_093 = find_row_idx(uc_table, 'UC-093')
row_093 = uc_table.rows[idx_093]
set_cell_text(row_093.cells[1], 'Xem tồn kho hiện tại (cảnh báo tồn tối thiểu: xem JOB-04, chưa triển khai)')

# UC-094: annotate
idx_094 = find_row_idx(uc_table, 'UC-094')
row_094 = uc_table.rows[idx_094]
set_cell_text(row_094.cells[1], 'Kiểm kê kho (chưa triển khai — hiện dựa vào chế độ điều chỉnh tồn kho gốc của Odoo)')

# UC-095: annotate
idx_095 = find_row_idx(uc_table, 'UC-095')
row_095 = uc_table.rows[idx_095]
set_cell_text(row_095.cells[1], 'Kiểm tra khả dụng vật tư theo BOM (chưa triển khai — backlog Phase 3, ATP)')

# UC-096: annotate
idx_096 = find_row_idx(uc_table, 'UC-096')
row_096 = uc_table.rows[idx_096]
set_cell_text(row_096.cells[1], 'Nhập kho thành phẩm theo BOM (chưa triển khai — chỉ mới seed sẵn loại phiếu, chưa có logic/màn hình)')

# UC-097: annotate
idx_097 = find_row_idx(uc_table, 'UC-097')
row_097 = uc_table.rows[idx_097]
set_cell_text(row_097.cells[1], 'Tính nhu cầu vật tư cho phiếu nhập kho theo BOM (chưa triển khai)')

# UC-098: remove row entirely (no such approval step exists in code)
idx_098 = find_row_idx(uc_table, 'UC-098')
tr_098 = uc_table.rows[idx_098]._tr
tr_098.getparent().remove(tr_098)

print('UC table edits done. Row count now:', len(uc_table.rows))

# ---------- 2. Insert new FT table after "Module: Kho & Mua hàng (INV)" heading ----------
body = d.element.body
heading_p = None
for child in body.iterchildren():
    if child.tag == qn('w:p'):
        p = Paragraph(child, d)
        if p.text.strip() == 'g. Module: Kho & Mua hàng (INV)':
            heading_p = p
            break
assert heading_p is not None, 'heading not found'

# Build the new table by appending to the doc (python-docx only supports append),
# then relocate its XML element right after the heading paragraph.
ft_rows = [
    ['FT-ID', 'Feature', 'Priority', 'Description', 'UC-ID', 'BF-ID', 'SOURCE'],
    ['FT-75', 'Thiết lập layout kho', 'Must',
     'Hệ thống dựng sẵn 1 kho, 3 khu vực (Nhận hàng/Xưởng/Thành phẩm) và các loại phiếu kho khi cài đặt module, không có màn hình cấu hình riêng cho người dùng',
     '—', '', 'Tự phát triển'],
    ['FT-76', 'Nhận hàng NCC 2 bước', 'Must',
     'Ghi nhận hàng về từ nhà cung cấp vào khu vực chờ kiểm; xác nhận phiếu tự sinh phiếu Kiểm & cất hàng kế tiếp',
     'UC-091', '', 'Tự phát triển (trên nền stock của Odoo)'],
    ['FT-77', 'Kiểm tra chất lượng hàng nhận (QC)', 'Must',
     'Nhập số lượng đạt/loại cho từng dòng hàng khi cất kho; bắt buộc lý do khi có hàng loại; chặn xác nhận nếu số liệu không hợp lệ hoặc thiếu lô',
     'UC-091A', '', 'Tự phát triển'],
    ['FT-78', 'Tự động tách và tạo phiếu trả hàng NCC', 'Must',
     'Khi có hàng bị loại ở bước kiểm, hệ thống tự tách phần hàng lỗi và tạo phiếu trả nhà cung cấp ở trạng thái nháp, giao cho bộ phận Mua hàng xử lý',
     'UC-091B', '', 'Tự phát triển'],
    ['FT-79', 'Tự động cấp số lô', 'Should',
     'Với vật tư/bán thành phẩm theo dõi theo lô, hệ thống tự sinh số lô theo mẫu khi nhận hàng nếu người dùng chưa nhập; vẫn cho phép sửa',
     'UC-091, UC-091A', '', 'Tự phát triển (trên nền lô của Odoo)'],
    ['FT-80', 'Truy vết lô hàng', 'Should',
     'Mỗi lô lưu lại nhà cung cấp, ngày nhận và phiếu nhập gốc; xem lại được từ màn Lô hàng',
     'UC-091', '', 'Tự phát triển'],
    ['FT-81', 'Xem tồn kho hiện tại', 'Must',
     'Danh sách tồn kho theo sản phẩm/lô/vị trí/nhà cung cấp, chỉ xem, không hiển thị giá vốn',
     'UC-093', '', 'Tự phát triển (trên nền tồn kho của Odoo)'],
    ['FT-83', 'Ghi nhận xuất kho', 'Must',
     'Chưa triển khai — dự kiến K6-K8 (giao hàng khách, xuất huỷ, xuất sản xuất thủ công)',
     'UC-092', '', '—'],
    ['FT-84', 'Kiểm kê kho', 'Should',
     'Chưa triển khai — hiện dựa vào chế độ điều chỉnh tồn kho gốc của Odoo, không có màn hình riêng của DLM-ERP',
     'UC-094', '', '—'],
]

new_table = d.add_table(rows=len(ft_rows), cols=7)
# copy style from an existing FT table (e.g. table index 5, AUTH module) for visual consistency
ref_table = d.tables[5]
new_table.style = ref_table.style
for r, values in enumerate(ft_rows):
    for c, val in enumerate(values):
        set_cell_text(new_table.rows[r].cells[c], val)

# relocate: move new_table._tbl right after heading_p._p
tbl_elm = new_table._tbl
tbl_elm.getparent().remove(tbl_elm)
heading_p._p.addnext(tbl_elm)

d.save(OUT)
print('Saved', OUT)
