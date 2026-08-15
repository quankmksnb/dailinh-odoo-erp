# -*- coding: utf-8 -*-
import copy
import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
from uc_index_data import UC_DATA

SRC = 'scratch_docx/edited8.docx'
OUT = 'scratch_docx/edited9.docx'

d = docx.Document(SRC)


def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)


ACTOR_NAMES = [
    'Quản trị hệ thống', 'CEO', 'Trưởng phòng Kinh doanh', 'Nhân viên Kinh doanh',
    'Kỹ thuật', 'Kế toán nội bộ / Thủ kho', 'Mua hàng',
]

REPORT_KEYWORDS = ['Xem ', 'Xem/', 'Tra cứu', 'Theo dõi', 'Giám sát', 'Tổng quan', 'Bảng tổng quan', 'Danh sách']
SYSTEM_ASSISTED = {
    'UC-010', 'UC-057', 'UC-067', 'UC-069A', 'UC-078', 'UC-079', 'UC-080',
    'UC-091', 'UC-107', 'UC-069',
}
FORCE_PLANNED = {'UC-114', 'UC-115', 'UC-116', 'UC-117', 'UC-118'}

uc_table = d.tables[2]
rows_data = []
for row in uc_table.rows[1:]:
    cells = [c.text.strip() for c in row.cells]
    uc_id, name = cells[0], cells[1]
    perms = cells[2:]

    if uc_id in FORCE_PLANNED:
        typ = 'Planned'
    elif all(p == '—' for p in perms):
        typ = 'System'
    elif '(Planned' in name:
        typ = 'Planned'
    elif any(kw in name for kw in REPORT_KEYWORDS):
        typ = 'UI / Report'
    elif uc_id in SYSTEM_ASSISTED:
        typ = 'UI / System-assisted'
    else:
        typ = 'UI'

    # thứ tự ưu tiên khi chọn actor chính: người trực tiếp thao tác trước,
    # người cấu hình/giám sát (CEO, SysAdmin) sau — tránh lệch về SysAdmin
    # chỉ vì cột đó đứng đầu bảng.
    PICK_ORDER = [4, 5, 3, 2, 6, 1, 0]

    if typ in ('System', 'Planned'):
        primary_actor = 'Hệ thống' if typ == 'System' else '—'
    elif all(p == 'Full' for p in perms):
        primary_actor = 'Tất cả người dùng'
    else:
        primary_actor = None
        for i in PICK_ORDER:
            if perms[i] == 'Full':
                primary_actor = ACTOR_NAMES[i]
                break
        if primary_actor is None:
            for i in PICK_ORDER:
                if perms[i] == 'Restricted':
                    primary_actor = ACTOR_NAMES[i]
                    break
        if primary_actor is None:
            primary_actor = '—'

    desc, bf, rel = UC_DATA.get(uc_id, ('', '', ''))
    rows_data.append([uc_id, name, typ, primary_actor, desc, bf, rel])

print('Total UC rows:', len(rows_data))

# find existing simplified index table (UC-ID | Use Case | Type, 3 cols) and replace it
target_elm = None
for t in d.tables:
    if t.rows[0].cells[0].text.strip() == 'UC-ID' and len(t.columns) == 3:
        target_elm = t._tbl
        ref_style = t.style
        break
assert target_elm is not None

header = ['UC-ID', 'Name', 'Type', 'Primary Actor', 'Description', 'BF-ID', 'Relationship']
new_table = d.add_table(rows=len(rows_data) + 1, cols=7)
new_table.style = ref_style
for c, val in enumerate(header):
    set_cell_text(new_table.rows[0].cells[c], val)
for r, values in enumerate(rows_data, start=1):
    for c, val in enumerate(values):
        set_cell_text(new_table.rows[r].cells[c], val)

new_tbl_elm = new_table._tbl
new_tbl_elm.getparent().remove(new_tbl_elm)
target_elm.addprevious(new_tbl_elm)
target_elm.getparent().remove(target_elm)

print('UC Index table replaced with full 7-column version.')

d.save(OUT)
print('Saved', OUT)
