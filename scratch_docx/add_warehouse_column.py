# -*- coding: utf-8 -*-
import copy
import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

SRC = 'scratch_docx/edited9.docx'
OUT = 'scratch_docx/edited10.docx'

d = docx.Document(SRC)


def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)


def find_row_idx(table, key, col=0):
    for i, row in enumerate(table.rows):
        if row.cells[col].text.strip() == key:
            return i
    raise ValueError(key)


# ============ 1. Add "Warehouse Keeper" column to Permission Matrix ============
uc_table = d.tables[2]
tbl_elm = uc_table._tbl

# extend tblGrid with one more gridCol (copy width of last column)
tblGrid = tbl_elm.find(qn('w:tblGrid'))
gridCols = tblGrid.findall(qn('w:gridCol'))
last_grid_col = gridCols[-1]
new_grid_col = copy.deepcopy(last_grid_col)
tblGrid.append(new_grid_col)

# for every row, clone the last <w:tc> and append
for tr in tbl_elm.findall(qn('w:tr')):
    tcs = tr.findall(qn('w:tc'))
    last_tc = tcs[-1]
    new_tc = copy.deepcopy(last_tc)
    tr.append(new_tc)

print('Column added. New column count:', len(uc_table.columns))

# header
header_row = uc_table.rows[0]
set_cell_text(header_row.cells[9], 'Warehouse Keeper')

# default all data rows to "No"
for row in uc_table.rows[1:]:
    set_cell_text(row.cells[9], 'No')

# ---- Kho UCs: set Warehouse Keeper values + adjust Internal Accountant (col 7) down where ACL is read-only ----
# perms tuple: (warehouse_val, accountant_new_val_or_None_to_keep)
KHO_PERMS = {
    'UC-090A': ('Full', 'No'),
    'UC-091':  ('Full', 'Restricted'),
    'UC-091A': ('Full', 'Restricted'),
    'UC-091B': ('Restricted', 'No'),
    'UC-092':  ('Full', 'Restricted'),
    'UC-092A': ('Full', 'Restricted'),
    'UC-092B': ('Full', 'Restricted'),
    'UC-092C': ('Full', 'Full'),   # ACL: dl.scrap.recovery.report đọc được cho cả Warehouse lẫn Accountant riêng
    'UC-093':  ('Full', None),     # giữ nguyên Accountant=Full (UC dạng xem/report, ACL đều có quyền đọc)
    'UC-094':  ('Full', 'No'),
}

for uc_id, (wh_val, acc_val) in KHO_PERMS.items():
    row = uc_table.rows[find_row_idx(uc_table, uc_id)]
    set_cell_text(row.cells[9], wh_val)
    if acc_val is not None:
        set_cell_text(row.cells[7], acc_val)

print('Kho UC permissions set for Warehouse Keeper column.')

d.save(OUT)
print('Saved', OUT)
