# -*- coding: utf-8 -*-
import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

SRC = 'scratch_docx/edited3b.docx'
OUT = 'scratch_docx/edited4.docx'

d = docx.Document(SRC)


def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)


def find_row_idx(table, key, col=0):
    for i, row in enumerate(table.rows):
        if row.cells[col].text.strip() == key:
            return i
    raise ValueError(key)


def find_heading(text):
    body = d.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            p = Paragraph(child, d)
            if p.text.strip() == text:
                return p
    raise ValueError('heading not found: ' + text)


# ============ 1. Permission matrix: UC-077, UC-102 ============
uc_table = d.tables[2]

row = uc_table.rows[find_row_idx(uc_table, 'UC-077')]
set_cell_text(row.cells[1], 'Mở khu vực CRM & Báo giá (menu Báo giá / Phê duyệt / Đơn bán hàng)')

row = uc_table.rows[find_row_idx(uc_table, 'UC-102')]
set_cell_text(row.cells[1], 'Cảnh báo cấu hình tham số giá bất hợp lý (chưa triển khai đúng mô tả — chỉ cảnh báo thiếu cấu hình gây lỗi báo giá, xem FT-17)')

print('UC-077, UC-102 updated.')

# ============ 2. TECH FT table: fill FT-64, enrich FT-68, add FT-97 ============
tech_heading = find_heading('f. Module: Kỹ thuật (TECH)')
tech_ft_table = Table(tech_heading._p.getnext(), d)
ref_style = tech_ft_table.style

row = tech_ft_table.rows[find_row_idx(tech_ft_table, 'FT-64')]
set_cell_text(row.cells[1], 'Tự tính định mức theo quy cách vật tư')
set_cell_text(row.cells[3], 'Với vật tư dạng cắt theo chiều dài hoặc cắt tấm, hệ thống tự tính số lượng cần dùng từ kích thước (dài/rộng, số lượng chiếc) nhập vào dòng định mức, quy đổi theo quy cách của vật tư (khổ tấm, chiều dài cây, khối lượng/mét...). Chưa hỗ trợ các hình dạng phức tạp hơn (tròn, ống, khối) — cơ chế "hình dạng đo lường" tổng quát trong thiết kế trước đã ngừng dùng.')

row = tech_ft_table.rows[find_row_idx(tech_ft_table, 'FT-68')]
set_cell_text(row.cells[3], 'Mỗi định mức có thể lưu nhiều phiên bản và hệ thống đánh dấu rõ phiên bản nào đang được dùng. Lưu ý: định mức có 2 loại mang ý nghĩa "phiên bản" khác nhau — định mức chuẩn của sản phẩm (tham gia đánh dấu "hiện hành") và định mức sinh riêng cho một báo giá/đơn hàng cụ thể (không tham gia đánh dấu "hiện hành", tồn tại song song).')

new_row_data = ['FT-97', 'Gợi ý sản phẩm đã từng gia công tương tự', 'Could',
                'Khi Kỹ thuật xử lý một dòng RFQ, hệ thống tự chấm điểm và gợi ý các sản phẩm đã có sẵn dựa trên nhiều tín hiệu (trùng/gần giống tên, cùng nhóm, cùng khách hàng, khớp kích thước, thuộc cùng họ BOM mẫu tham số...); tự chọn sẵn khi điểm đủ cao, chỉ gợi ý khi điểm ở mức trung bình',
                'UC-067', 'BF-02', 'Tự phát triển']
last_row_tr = tech_ft_table.rows[-1]._tr
import copy as _copy
new_tr = _copy.deepcopy(last_row_tr)
last_row_tr.addnext(new_tr)
new_row = None
for r in tech_ft_table.rows:
    if r._tr is new_tr:
        new_row = r
        break
for cell, text in zip(new_row.cells, new_row_data):
    set_cell_text(cell, text)

print('TECH FT table updated (FT-64 filled, FT-68 enriched, FT-97 added).')

# ============ 3. SCR table fixes ============
scr_table = None
for t in d.tables:
    if t.rows[0].cells[0].text.strip() == 'SCR-ID':
        scr_table = t
        break
assert scr_table is not None

fixes = {
    'SCR-17': 'UC-067, UC-068, UC-069',
    'SCR-18': 'UC-073',
    'SCR-19': 'UC-075',
    'SCR-20': '—',
    'SCR-21': 'UC-079, UC-080, UC-083, UC-084, UC-086, UC-087A, UC-088, UC-107, UC-108',
    'SCR-22': 'UC-086, UC-087, UC-087A',
    'SCR-24': 'UC-089',
    'SCR-25': 'UC-089',
    'SCR-28': 'UC-084',
}
for scr_id, new_uc in fixes.items():
    row = scr_table.rows[find_row_idx(scr_table, scr_id)]
    set_cell_text(row.cells[5], new_uc)

print('SCR table fixed:', list(fixes.keys()))

d.save(OUT)
print('Saved', OUT)
