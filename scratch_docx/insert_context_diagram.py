# -*- coding: utf-8 -*-
import docx
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = 'scratch_docx/edited11.docx'
OUT = 'scratch_docx/edited12.docx'
IMG = 'scratch_docx/context_diagram.png'

d = docx.Document(SRC)

body = d.element.body
sectPr = body.find(qn('w:sectPr'))
assert sectPr is not None

# Build new content as a temporary set of paragraphs appended at document end,
# then move them (as XML elements) to right before sectPr — consistent with
# the approach used for tables earlier in this session.
heading = d.add_paragraph('4.5 Context Diagram (Sơ đồ ngữ cảnh hệ thống)')
heading.style = d.styles['Heading 2']

intro = d.add_paragraph(
    'Sơ đồ dưới đây thể hiện DLM-ERP như một khối xử lý trung tâm, cùng các tác nhân/hệ '
    'thống bên ngoài biên hệ thống và luồng dữ liệu ra/vào giữa chúng. Đường liền nét màu '
    'xanh là luồng dữ liệu qua giao diện hệ thống thật (người dùng thao tác, hoặc hệ thống '
    'tự gửi email); đường đứt nét là mối quan hệ không có kết nối điện tử trực tiếp hoặc nằm '
    'ngoài biên hệ thống.'
)

pic_p = d.add_paragraph()
pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pic_p.add_run()
run.add_picture(IMG, width=Inches(6.3))

caption = d.add_paragraph('Hình: Sơ đồ ngữ cảnh hệ thống (Context Diagram) — DLM-ERP')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
if caption.runs:
    caption.runs[0].italic = True

new_elms = [heading._p, intro._p, pic_p._p, caption._p]
for elm in new_elms:
    elm.getparent().remove(elm)
for elm in new_elms:
    sectPr.addprevious(elm)

print('Inserted context diagram section before sectPr.')

d.save(OUT)
print('Saved', OUT)
