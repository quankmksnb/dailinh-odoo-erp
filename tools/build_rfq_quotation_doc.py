from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Dac_ta_luong_tao_bao_gia_tu_RFQ_BOM_v1.1.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FB"
PALE_YELLOW = "FFF4CE"
PALE_RED = "FCE8E6"
PALE_GREEN = "E7F4E4"
GRAY = "666666"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = round(sum(widths_in) * 1440)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_in:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(round(w * 1440)))
        grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            w = round(widths_in[i] * 1440)
            cell.width = Inches(widths_in[i])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(w))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=None, bold=None, color=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for sname, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        s = styles[sname]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    for sname in ("List Bullet", "List Number"):
        s = styles[sname]
        s.font.name = "Calibri"
        s.font.size = Pt(11)
        s.paragraph_format.left_indent = Inches(0.375)
        s.paragraph_format.first_line_indent = Inches(-0.188)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.25
    if "Callout" not in styles:
        s = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        s = styles["Callout"]
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)
    s.paragraph_format.left_indent = Inches(0.15)
    s.paragraph_format.right_indent = Inches(0.15)
    s.paragraph_format.space_before = Pt(5)
    s.paragraph_format.space_after = Pt(7)
    s.paragraph_format.line_spacing = 1.15


def add_run_parts(p, parts):
    for text, bold, color, italic in parts:
        r = p.add_run(text)
        set_font(r, bold=bold, color=color, italic=italic)


def add_callout(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.style = doc.styles["Callout"]
    r = p.add_run(label + " ")
    set_font(r, bold=True, color=DARK_BLUE)
    set_font(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, font_size=9.2, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_shading(hdr.cells[i], header_fill)
        p = hdr.cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=font_size, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(val))
            set_font(r, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def number(doc, text):
    return doc.add_paragraph(text, style="List Number")


def field_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run("Đặc tả RFQ → BOM → Báo giá  |  Trang ")
    set_font(r, size=8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    sec = doc.sections[0]
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = header.add_run("DLM ERP  |  Tài liệu phân tích & đặc tả triển khai")
    set_font(rr, size=8.5, color=GRAY)
    field_footer(sec)

    # Editorial-cover opening.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("ĐẶC TẢ NGHIỆP VỤ VÀ KỸ THUẬT"), size=11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("TẠO BÁO GIÁ TỪ RFQ VÀ BOM"), size=27, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_font(p.add_run("Sản phẩm thương mại, sản phẩm gia công, giá vật tư và cấu hình tính giá"), size=14, color=GRAY)
    add_callout(doc, "Kết luận rà soát:", "Code hiện tại đã xử lý RFQ, BOM và giá vật tư cơ bản, nhưng chưa có pricing engine nối các model cấu hình vào dl.quotation. Tài liệu này mô tả chính xác hiện trạng và đặc tả phần cần bổ sung để tạo báo giá tự động, có thể truy vết và phê duyệt.", PALE_YELLOW)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    set_font(p.add_run("Phiên bản tài liệu: 1.1  |  Cập nhật Decision log & kế hoạch P0: 16/07/2026"), size=10, color=GRAY)
    doc.add_page_break()

    doc.add_heading("1. Mục tiêu và phạm vi", level=1)
    doc.add_paragraph("Mục tiêu là tạo một báo giá bán hàng từ một Yêu cầu báo giá (RFQ) sau khi Kỹ thuật đã xử lý các dòng gia công bằng BOM. Báo giá phải phân biệt rõ hàng thương mại và hàng gia công, lấy đúng giá vật tư, áp dụng đúng các cấu hình đang hiệu lực, lưu snapshot để báo giá cũ không đổi khi cấu hình thay đổi, và định tuyến phê duyệt theo giá trị/chiết khấu/giá sàn.")
    doc.add_heading("1.1. Kết quả đầu ra mong muốn", level=2)
    for t in [
        "Mỗi RFQ đã xử lý xong tạo được một dl.quotation và các dl.quotation.line tương ứng.",
        "Dòng thương mại lấy giá bán có sẵn của sản phẩm; không tạo và không tính BOM.",
        "Dòng gia công lấy BOM kỹ thuật đã xác nhận/khóa, tính giá vật tư và các lớp chi phí cấu hình.",
        "Báo giá thể hiện được giá trước chiết khấu, chiết khấu, giá sau chiết khấu, VAT, tổng thanh toán và trạng thái phê duyệt.",
        "Mỗi con số có thể truy ngược về BOM, bảng giá nhà cung cấp và revision cấu hình đã sử dụng.",
    ]:
        bullet(doc, t)
    doc.add_heading("1.2. Nguyên tắc nghiệp vụ chốt", level=2)
    add_table(doc, ["Loại dòng RFQ", "Nguồn sản phẩm", "Nguồn giá chính", "BOM"], [
        ["Thương mại", "resolved_product_id (product_kind=trading)", "list_price theo code hiện tại", "Không dùng"],
        ["Gia công", "resolved_product_id (manufactured)", "Chi phí BOM + công đoạn + chi phí chung + markup", "Bắt buộc confirmed/locked"],
    ], [1.15, 2.05, 2.25, 1.05])
    add_callout(doc, "Phạm vi giá thương mại:", "Theo yêu cầu 'giá sẵn cứ vậy mà bán', đơn giá gốc của dòng thương mại được lấy trực tiếp từ list_price. Chiết khấu/VAT ở cấp báo giá vẫn có thể áp dụng nếu nghiệp vụ mong muốn; BOM và các chi phí gia công không áp dụng cho dòng này.")

    doc.add_heading("2. Hiện trạng code", level=1)
    doc.add_heading("2.1. Những phần đã có và có thể tái sử dụng", level=2)
    add_table(doc, ["Khối", "Model/field chính", "Hành vi hiện tại"], [
        ["RFQ", "dl.quotation.request / line_ids", "Theo dõi new → processing → confirmed → quoted; mỗi dòng có product_type, quantity, resolved_product_id và resolved_bom_id."],
        ["Xử lý kỹ thuật", "dl.rfq.resolve.wizard", "Cho phép chọn/tạo sản phẩm gia công, tạo/chọn BOM báo giá và ghi kết quả về dòng RFQ."],
        ["BOM", "dl.bom / dl.bom.line", "Tổng chi phí vật tư = tổng subtotal các dòng BOM."],
        ["Giá vật tư", "product.supplierinfo", "Chỉ bảng giá approved và is_applied mới được dùng làm price_snapshot."],
        ["Hao hụt BOM", "effective_qty", "quantity × (1 + waste_rate/100)."],
        ["Cấu hình", "dl.pricing.*", "Đã có model cho hao hụt, công đoạn, chi phí chung, markup, chiết khấu và ma trận phê duyệt."],
        ["Báo giá", "dl.quotation / dl.quotation.line", "Hiện chỉ có qty, price_unit và subtotal = qty × price_unit."],
    ], [1.05, 2.05, 3.4])
    doc.add_heading("2.2. Khoảng trống quan trọng", level=2)
    for t in [
        "Chưa có quan hệ từ dl.quotation về RFQ (quotation_request_id) và từ dòng báo giá về RFQ line/product/BOM.",
        "Chưa có action_create_quotation; action_mark_quoted chỉ đổi trạng thái và không tạo bản ghi báo giá.",
        "dl.quotation chưa đọc dl.pricing.operation.rule, dl.pricing.cost.adjustment.rule, dl.pricing.profit.rule, dl.pricing.discount.rule hoặc approval matrix.",
        "Các rule có revision/validity/used_in_snapshot nhưng chưa có model snapshot trên báo giá để thực sự bảo toàn lịch sử.",
        "Có hai cơ chế hao hụt: dl.pricing.waste cũ được BOM dùng trực tiếp và dl.pricing.waste.rule mới có category/product, recovery, complexity nhưng chưa được nối với BOM.",
        "BOM bán thành phẩm đang lấy toàn bộ total_material_cost của BOM con làm đơn giá, chưa chia cho product_qty; điều này sai nếu BOM con tạo ra nhiều hơn 1 đơn vị.",
        "price_snapshot là compute stored nhưng phụ thuộc trực tiếp seller_ids; khi đổi bảng giá áp dụng, BOM cũ có thể bị tính lại, chưa đúng nghĩa snapshot bất biến.",
        "Chưa kiểm tra ngày hiệu lực bảng giá NCC và quy đổi tiền tệ/UoM trong phép tính BOM.",
    ]:
        bullet(doc, t)
    add_callout(doc, "Đánh giá sẵn sàng:", "Nền dữ liệu đủ tốt để xây engine, nhưng chưa thể khẳng định hệ thống hiện tại tự tạo báo giá đúng cấu hình. Cần bổ sung lớp dịch vụ tính giá, trường snapshot, action tạo báo giá và kiểm thử.", PALE_RED)

    doc.add_heading("3. Luồng nghiệp vụ đầu-cuối", level=1)
    steps = [
        ("Sales tạo RFQ", "Nhập khách hàng, hạn yêu cầu và các dòng. Dòng thương mại phải chọn sản phẩm; dòng gia công nhập tên/yêu cầu/kích thước và có thể chọn sản phẩm tham khảo."),
        ("Kỹ thuật xử lý dòng gia công", "Chọn hoặc tạo sản phẩm manufactured; tạo/chọn BOM quotation; khai báo vật tư, định mức, hao hụt và công đoạn cần thiết."),
        ("Xác nhận BOM", "BOM dùng cho RFQ phải ở confirmed hoặc locked. Mọi vật tư thô phải có đúng một supplierinfo đã duyệt và đang áp dụng."),
        ("RFQ chuyển confirmed", "Tất cả dòng phải khả thi và đã resolve: thương mại cần product; gia công cần product + BOM."),
        ("Sales bấm Tạo báo giá", "Hệ thống khóa logic theo một ngày tính giá, tải các cấu hình active đúng company/validity, tính từng dòng và lưu snapshot."),
        ("Đánh giá phê duyệt", "So tổng sau chiết khấu trước VAT, mức chiết khấu và giá sàn; chọn cấp cao nhất nếu nhiều điều kiện đồng thời phát sinh."),
        ("Gửi/duyệt báo giá", "Không cần duyệt thì Ready to Send; cần duyệt thì tạo một approval request. Sau duyệt mới được gửi khách."),
        ("Đóng RFQ", "Khi báo giá tạo thành công, ghi quotation_id và đổi RFQ sang quoted trong cùng transaction."),
    ]
    add_table(doc, ["Bước", "Chủ trì", "Xử lý"], [[i+1, ["Sales","Kỹ thuật","Kỹ thuật","Hệ thống","Sales/Hệ thống","Hệ thống","Sales/Quản lý","Hệ thống"][i], f"{a}: {b}"] for i,(a,b) in enumerate(steps)], [0.55, 1.15, 4.8], 9)

    doc.add_heading("4. Điều kiện trước khi tạo báo giá", level=1)
    doc.add_heading("4.1. Điều kiện chung của RFQ", level=2)
    for t in [
        "RFQ.status = confirmed và chưa có báo giá hiệu lực được tạo từ RFQ này.",
        "Khách hàng hợp lệ, có currency/company rõ ràng; quantity của mọi dòng > 0.",
        "Không đưa dòng is_infeasible vào báo giá. Nếu RFQ có dòng không khả thi, cần quy định rõ: bỏ dòng có cảnh báo hoặc chặn toàn RFQ; khuyến nghị chặn để Sales xác nhận lại phạm vi.",
    ]: bullet(doc, t)
    doc.add_heading("4.2. Dòng thương mại", level=2)
    for t in [
        "product_type = trading; resolved_product_id.product_kind = trading.",
        "Không có resolved_bom_id; nếu có phải báo lỗi dữ liệu.",
        "list_price phải > 0. Nếu bằng 0, không tự tạo báo giá vì sẽ phát sinh dòng giá 0 ngoài ý muốn.",
    ]: bullet(doc, t)
    doc.add_heading("4.3. Dòng gia công", level=2)
    for t in [
        "resolved_product_id thuộc manufactured (hoặc phạm vi được duyệt rõ ràng); resolved_bom_id thuộc đúng sản phẩm.",
        "BOM.status thuộc confirmed/locked và bom_type = quotation.",
        "BOM.product_qty > 0; BOM có ít nhất một dòng vật tư.",
        "Mỗi dòng vật tư có effective_qty > 0 và price_snapshot > 0; nếu là material_processed phải tìm thấy BOM con confirmed/locked.",
        "Không tồn tại vòng lặp BOM (A dùng B, B dùng lại A).",
    ]: bullet(doc, t)

    doc.add_heading("5. Nguồn giá vật tư và cách tính BOM", level=1)
    doc.add_heading("5.1. Vật tư thô", level=2)
    doc.add_paragraph("Với component.product_kind = material, hệ thống chọn product.supplierinfo có is_applied=True. Constraint hiện tại bảo đảm bảng giá này đã approved và mỗi product template chỉ có tối đa một bảng giá đang áp dụng.")
    add_callout(doc, "Công thức dòng BOM:", "effective_qty = quantity × (1 + waste_rate/100); material_subtotal = effective_qty × unit_material_price.")
    doc.add_heading("5.2. Vật tư đã gia công/bán thành phẩm", level=2)
    doc.add_paragraph("Với material_processed, chi phí phải được tính đệ quy từ BOM con mới nhất đang confirmed/locked. Đơn giá đúng của bán thành phẩm phải quy về một đơn vị đầu ra:")
    add_callout(doc, "Công thức cần dùng:", "processed_unit_cost = child_bom.total_cost / child_bom.product_qty. Không dùng trực tiếp child_bom.total_material_cost làm đơn giá nếu product_qty có thể khác 1.", PALE_YELLOW)
    doc.add_heading("5.3. Hao hụt và thu hồi", level=2)
    doc.add_paragraph("Hiện BOM chỉ dùng waste_rate đơn giản. Khi nối pricing_waste.rule, quy tắc vật tư cụ thể ưu tiên hơn category; hệ số phức tạp nhân vào tỷ lệ hao hụt cơ sở. Nếu có thu hồi phế liệu, giá trị thu hồi phải trừ khỏi chi phí vật tư, không trừ khỏi số lượng vật tư mua.")
    add_table(doc, ["Thành phần", "Công thức đề xuất"], [
        ["Tỷ lệ hao hụt áp dụng", "applied_waste_pct = base_waste_pct × complexity_factor"],
        ["Lượng hao hụt", "waste_qty = net_qty × applied_waste_pct / 100"],
        ["Lượng tính giá", "gross_qty = net_qty + waste_qty"],
        ["Lượng phế thu hồi", "recovered_qty = waste_qty × recovery_rate / 100"],
        ["Giá trị thu hồi", "recovery_value = recovered_qty × scrap_unit_price"],
        ["Chi phí vật tư ròng", "gross_qty × material_unit_price − recovery_value"],
    ], [2.0, 4.5])

    doc.add_heading("6. Engine tính giá cho dòng gia công", level=1)
    doc.add_heading("6.1. Thứ tự tính bắt buộc", level=2)
    stages = [
        ["A", "Material cost", "Tổng chi phí vật tư ròng từ BOM, quy về một đơn vị đầu ra."],
        ["B", "Operation cost", "Tổng chi phí các công đoạn active, gồm setup fee theo lô."],
        ["C", "Direct cost", "A + B."],
        ["D", "Cost adjustments", "Chi phí chung/phụ phí/hệ số theo đúng method và condition."],
        ["E", "Total cost", "Direct cost + các khoản điều chỉnh; factor áp vào cơ sở được định nghĩa."],
        ["F", "Target selling price", "Total cost × (1 + target_markup/100)."],
        ["G", "Floor price", "Total cost × (1 + min_markup/100)."],
        ["H", "Rounding", "Làm tròn giá bán theo rounding_to sau khi cộng markup, trước chiết khấu."],
    ]
    add_table(doc, ["Lớp", "Tên", "Ý nghĩa"], stages, [0.55, 1.65, 4.3])
    doc.add_heading("6.2. Công đoạn", level=2)
    add_table(doc, ["method", "Công thức chi phí công đoạn"], [
        ["percent_material", "related_material_cost × price_rate/100 + setup_fee_per_batch"],
        ["per_kg", "weight_kg × price_rate + setup_fee_per_batch"],
        ["per_meter", "length_m × price_rate + setup_fee_per_batch"],
        ["per_sqm", "area_m² × price_rate + setup_fee_per_batch"],
        ["per_unit", "output_qty × price_rate + setup_fee_per_batch"],
        ["per_batch", "price_rate + setup_fee_per_batch (chỉ một lần cho lô)"],
    ], [1.55, 4.95])
    add_callout(doc, "Điểm phải bổ sung vào dữ liệu:", "BOM hiện chưa có danh sách công đoạn được chọn. Cần model trung gian (ví dụ dl.bom.operation.line) chứa operation_rule_id, base_qty/base_amount và snapshot đơn giá để biết công đoạn nào áp dụng cho sản phẩm cụ thể.", PALE_YELLOW)
    doc.add_heading("6.3. Chi phí chung và hệ số điều chỉnh", level=2)
    add_table(doc, ["method", "Cơ sở/công thức đề xuất"], [
        ["percent_direct", "direct_cost × value/100"],
        ["percent_cost", "running_cost × value/100; cần chốt thứ tự sequence để tránh vòng lặp"],
        ["per_unit", "quotation_line_qty × value"],
        ["per_batch", "value một lần cho dòng/lô"],
        ["fixed", "value một lần"],
        ["factor", "running_cost × (value − 1); value=1 không làm thay đổi giá"],
    ], [1.55, 4.95])
    doc.add_paragraph("Điều kiện urgent áp dụng khi delivery_days ≤ condition_days. Điều kiện small_order áp dụng khi giá trị cơ sở đơn hàng ≤ condition_amount. Khoản no_discount phải tách khỏi discountable_amount và chỉ cộng lại sau chiết khấu.")

    doc.add_heading("7. Giá thương mại, markup, chiết khấu, VAT và làm tròn", level=1)
    doc.add_heading("7.1. Dòng thương mại", level=2)
    add_callout(doc, "Quy tắc:", "base_unit_price = resolved_product_id.list_price. Không đọc BOM, operation rule hay cost adjustment dành cho gia công. Cần snapshot list_price vào dòng báo giá tại thời điểm tạo.", PALE_GREEN)
    doc.add_heading("7.2. Dòng gia công", level=2)
    doc.add_paragraph("Markup là phần trăm cộng trên giá thành, không phải margin trên doanh thu. Vì model dùng target_markup/min_markup, công thức phải là cost × (1 + rate/100). Không dùng cost/(1−rate) trừ khi đổi hẳn ý nghĩa cấu hình sang gross margin.")
    doc.add_heading("7.3. Chiết khấu", level=2)
    doc.add_paragraph("Discount rule active được chọn theo partner.dlm_customer_group. default_rate tự điền; max_rate là ngưỡng ngoại lệ. Khuyến nghị hỗ trợ discount_pct ở header và có thể override theo line, nhưng phải snapshot tỷ lệ thực dùng.")
    add_callout(doc, "Phân bổ chiết khấu để kiểm tra giá sàn:", "Chiết khấu header được phân bổ pro-rata theo subtotal trước chiết khấu của từng dòng: allocated_discount_i = header_discount_amount × line_subtotal_i / discountable_amount. Với dòng gia công, net_line_amount = line_subtotal − allocated_discount_i và net_unit_price = net_line_amount / qty; nếu net_unit_price < floor_price thì below_floor=True. Khoản no_discount không tham gia mẫu số phân bổ.", PALE_GREEN)
    add_table(doc, ["Giá trị", "Công thức"], [
        ["discountable_amount", "Tổng các khoản chịu chiết khấu"],
        ["discount_amount", "discountable_amount × discount_pct/100"],
        ["amount_before_vat", "discountable_amount − discount_amount + no_discount_amount"],
        ["vat_amount", "amount_before_vat × vat_pct/100"],
        ["amount_total", "amount_before_vat + vat_amount"],
        ["effective_markup", "(amount_before_vat − total_cost) / total_cost × 100"],
    ], [2.1, 4.4])
    doc.add_heading("7.4. Làm tròn", level=2)
    doc.add_paragraph("rounding_to = 0 nghĩa là không làm tròn. Với 1.000 hoặc 10.000 đồng, cần chốt nguyên tắc làm tròn gần nhất (khuyến nghị ROUND_HALF_UP). Làm tròn giá đơn vị mục tiêu trước chiết khấu; phần VAT và tổng tiền dùng currency rounding của Odoo.")

    doc.add_heading("8. Phê duyệt báo giá", level=1)
    doc.add_paragraph("pricing_matrix.evaluate_quotation đã có logic kiểm tra đồng thời ba trục và chọn cấp có rank cao nhất. Giá trị đưa vào ma trận là tổng sau chiết khấu, trước VAT.")
    add_table(doc, ["Điều kiện", "Cấp tối thiểu theo code", "Ghi chú"], [
        ["Giá trị chạm ngưỡng active", "Theo dòng matrix", "Chọn value_from lớn nhất ≤ amount."],
        ["Chiết khấu > mặc định nhưng ≤ tối đa", "Trưởng kinh doanh", "discount_above_default=True."],
        ["Chiết khấu > tối đa", "CEO", "discount_above_max=True."],
        ["Giá sau chiết khấu < giá sàn", "CEO", "below_floor=True."],
    ], [2.25, 1.65, 2.6])
    doc.add_heading("8.1. Trạng thái đề xuất", level=2)
    doc.add_paragraph("draft → pricing_ready → pending_approval → approved → sent; nhánh rejected quay lại draft sau chỉnh sửa. Nếu không cần duyệt, pricing_ready có thể chuyển thẳng approved/ready_to_send theo chính sách.")
    doc.add_heading("8.2. Hủy duyệt khi thay đổi", level=2)
    doc.add_paragraph("Bất kỳ thay đổi nào tới quantity, BOM, giá snapshot, discount, VAT, adjustment hoặc delivery condition đều phải gọi action_cancel_on_change cho yêu cầu pending/approved cũ, tính lại toàn bộ và đánh giá lại matrix. Không tái sử dụng kết quả duyệt cho một cấu trúc giá khác.")

    doc.add_heading("9. Thiết kế dữ liệu cần bổ sung", level=1)
    add_table(doc, ["Model", "Field đề xuất", "Mục đích"], [
        ["dl.quotation", "quotation_request_id", "Truy vết RFQ nguồn; unique nếu mỗi RFQ chỉ có một báo giá hiện hành."],
        ["dl.quotation", "company_id, pricing_date", "Chọn đúng rule theo công ty/ngày hiệu lực."],
        ["dl.quotation", "amount_untaxed, discount_amount, amount_before_vat, vat_amount, amount_total", "Tách rõ các lớp tiền."],
        ["dl.quotation", "discount_pct, vat_pct, total_cost, floor_amount, effective_markup", "Đánh giá thương mại và phê duyệt."],
        ["dl.quotation", "approval_state/request_id", "Liên kết luồng duyệt."],
        ["dl.quotation.line", "rfq_line_id, product_id, bom_id, line_type", "Truy nguồn từng dòng."],
        ["dl.quotation.line", "material_cost, operation_cost, adjustment_cost, total_cost", "Bảng phân tích nội bộ."],
        ["dl.quotation.line", "base_price, floor_price, discount_pct, price_unit", "Giá mục tiêu, giá sàn và giá bán."],
        ["dl.quotation.price.component", "component_type, source_model/id/revision, base, rate, amount, no_discount", "Snapshot từng cấu phần giá và rule nguồn."],
        ["dl.bom.operation.line", "operation_rule_id, method, base_qty, price_rate, setup_fee, subtotal", "Chọn công đoạn cụ thể trên BOM."],
    ], [1.35, 2.4, 2.75], 8.6)
    add_callout(doc, "Bảo mật:", "Các trường giá thành/material cost chỉ cho Accountant, Sales Manager, CEO và Admin như BOM hiện tại. Sales thường chỉ cần xem giá bán, chiết khấu và lý do cần duyệt.")

    doc.add_heading("10. Thuật toán tạo báo giá đề xuất", level=1)
    algo = [
        "ensure_one(); kiểm tra RFQ confirmed, không cancelled/quoted và không có báo giá trùng.",
        "Xác định pricing_date, company, currency; tải profit rule, discount rule, waste/operation/cost rules active trong ngày.",
        "Tạo quotation header ở draft với quotation_request_id và snapshot các tham số header.",
        "Duyệt từng RFQ line: trading → snapshot list_price; manufactured → validate BOM rồi gọi pricing service tính breakdown.",
        "Tạo quotation line và price component snapshots. Không để compute động dựa vào rule active sau này.",
        "Tính tổng discountable/no_discount, chiết khấu, trước VAT, VAT và tổng thanh toán.",
        "Tính floor_amount và các cờ discount_above_default, discount_above_max, below_floor.",
        "Gọi approval matrix evaluate_quotation(amount_before_vat, ...). Nếu required, tạo approval request và chuyển pending_approval.",
        "Gắn quotation_id vào RFQ; gọi action_mark_quoted trong cùng transaction; trả action mở báo giá.",
        "Nếu bất kỳ bước nào lỗi, rollback toàn bộ: không để RFQ quoted nhưng thiếu báo giá hoặc báo giá dở dang.",
    ]
    for a in algo: number(doc, a)
    doc.add_heading("10.1. Pseudocode", level=2)
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    code = """quotation = create_header(rfq, pricing_context)\nfor rfq_line in rfq.line_ids:\n    if rfq_line.product_type == 'trading':\n        result = price_trading(rfq_line.resolved_product_id.list_price)\n    else:\n        result = pricing_service.price_bom(rfq_line.resolved_bom_id, rfq_line.quantity)\n    create_line_and_snapshots(quotation, rfq_line, result)\nquotation.recompute_totals()\nevaluation = matrix.evaluate_quotation(quotation.amount_before_vat, flags...)\nquotation.apply_approval_result(evaluation)\nrfq.write({'quotation_id': quotation.id, 'status': 'quoted'})"""
    r = p.add_run(code)
    set_font(r, size=9, name="Consolas")

    doc.add_heading("11. Ví dụ tính giá minh họa", level=1)
    doc.add_paragraph("RFQ gồm 2 dòng: (1) hàng thương mại số lượng 2, list_price 1.500.000đ; (2) hàng gia công số lượng 10. Ví dụ chỉ nhằm minh họa thứ tự tính, không phải dữ liệu seed mặc định.")
    add_table(doc, ["Khoản", "Dòng thương mại", "Dòng gia công"], [
        ["Giá gốc", "2 × 1.500.000 = 3.000.000", "Vật tư ròng/đv = 600.000"],
        ["Công đoạn", "Không áp dụng", "150.000/đv"],
        ["Chi phí chung", "Không áp dụng", "75.000/đv"],
        ["Giá thành", "Không tính theo BOM", "825.000/đv"],
        ["Markup mục tiêu 12%", "Không cộng lại vì dùng giá sẵn", "99.000/đv"],
        ["Giá trước CK", "3.000.000", "10 × 924.000 = 9.240.000"],
        ["Cộng trước CK", "", "12.240.000"],
        ["Chiết khấu 5%", "", "612.000"],
        ["Sau CK, trước VAT", "", "11.628.000"],
        ["VAT 10%", "", "1.162.800"],
        ["Tổng thanh toán", "", "12.790.800"],
    ], [2.0, 2.1, 2.4])
    add_callout(doc, "Giá sàn dòng gia công:", "Nếu min_markup = 8%, floor_unit = 825.000 × 1,08 = 891.000đ. Giá sau chiết khấu quy về dòng phải được so với sàn; nếu thấp hơn, đánh cờ below_floor và chuyển CEO.", PALE_GREEN)
    doc.add_heading("11.1. Ví dụ riêng cho setup_fee theo lô", level=2)
    doc.add_paragraph("Một dòng gia công có số lượng 10, công đoạn Cắt có price_rate theo sản phẩm là 40.000đ/đơn vị và setup_fee là 300.000đ/lô. Quyết định B3 xác định mỗi dòng gia công là một lô độc lập.")
    add_table(doc, ["Khoản", "Cách tính", "Kết quả"], [
        ["Chi phí biến đổi", "10 × 40.000", "400.000đ/lô"],
        ["Setup", "Cộng đúng một lần cho dòng", "300.000đ/lô"],
        ["Tổng công đoạn", "400.000 + 300.000", "700.000đ/lô"],
        ["Phân bổ về đơn vị", "700.000 / 10", "70.000đ/đơn vị"],
    ], [1.65, 2.85, 2.0], 9.0)
    add_callout(doc, "Ràng buộc:", "Khi quantity thay đổi, setup_fee vẫn giữ nguyên cho lô nhưng phần setup phân bổ trên một đơn vị thay đổi. Không cộng setup_fee cho từng đơn vị.", PALE_YELLOW)

    doc.add_heading("12. Validation và thông báo lỗi", level=1)
    add_table(doc, ["Mã gợi ý", "Tình huống", "Thông báo người dùng"], [
        ["QTE-001", "RFQ chưa confirmed", "Chỉ RFQ đã xử lý xong mới được tạo báo giá."],
        ["QTE-002", "Dòng gia công thiếu BOM", "Sản phẩm gia công ... chưa có BOM đã xác nhận/khóa."],
        ["QTE-003", "Vật tư thiếu giá", "Vật tư ... chưa có bảng giá đã duyệt và đang áp dụng."],
        ["QTE-004", "BOM con thiếu/loop", "Không thể tính bán thành phẩm ...; kiểm tra BOM con hoặc vòng lặp."],
        ["QTE-005", "Thiếu profit rule", "Chưa có cấu hình lợi nhuận đang áp dụng tại ngày báo giá."],
        ["QTE-006", "Thiếu discount group/rule", "Khách hàng chưa có nhóm hoặc chưa có chính sách chiết khấu active."],
        ["QTE-007", "Dữ liệu đo/UoM không tương thích", "Không thể quy đổi định mức ... sang đơn vị giá vật tư."],
        ["QTE-008", "Báo giá trùng", "RFQ đã có báo giá; hãy mở báo giá hiện có hoặc tạo revision."],
    ], [0.8, 2.0, 3.7], 8.8)

    doc.add_heading("13. Kiểm thử chấp nhận", level=1)
    tests = [
        ["TC01", "RFQ chỉ có trading", "Giá dòng = list_price snapshot; không có BOM component."],
        ["TC02", "Gia công, vật tư thô có giá applied", "Subtotal BOM = effective_qty × price; tổng giá đúng."],
        ["TC03", "BOM có material_processed, product_qty > 1", "Đơn giá BOM con được chia product_qty."],
        ["TC04", "Hao hụt + complexity + recovery", "Chi phí ròng đúng theo công thức và truy được rule."],
        ["TC05", "6 method công đoạn", "Mỗi method dùng đúng base; setup chỉ cộng một lần/lô."],
        ["TC06", "6 method adjustment", "Đúng thứ tự và điều kiện urgent/small_order."],
        ["TC07", "Chiết khấu trên mặc định", "Tạo yêu cầu Trưởng KD nếu không có điều kiện cao hơn."],
        ["TC08", "Chiết khấu vượt max hoặc dưới sàn", "Yêu cầu CEO; reasons ghi đầy đủ."],
        ["TC09", "Đổi rule sau khi tạo", "Báo giá cũ không đổi; báo giá mới dùng revision mới."],
        ["TC10", "Đổi dữ liệu báo giá đã duyệt", "Duyệt cũ bị cancel; hệ thống tính/định tuyến lại."],
        ["TC11", "Một vật tư thiếu giá", "Transaction rollback; RFQ không chuyển quoted."],
        ["TC12", "Hai người tạo đồng thời", "Chỉ một báo giá được tạo; request còn lại nhận lỗi trùng."],
    ]
    add_table(doc, ["ID", "Kịch bản", "Kết quả mong đợi"], tests, [0.65, 2.35, 3.5], 8.8)

    doc.add_heading("14. Lộ trình triển khai", level=1)
    add_table(doc, ["Ưu tiên", "Hạng mục", "Kết quả"], [
        ["P0", "Bổ sung liên kết RFQ–Quotation, trường breakdown/totals và action_create_quotation", "Tạo được báo giá tự động cơ bản."],
        ["P0", "Pricing service cho trading + BOM material cost; sửa unit cost BOM con; khóa snapshot", "Giá nền đúng và ổn định."],
        ["P1", "BOM operation lines + operation rules", "Tính đủ chi phí công đoạn."],
        ["P1", "Cost adjustment rules + profit/discount + rounding/VAT", "Tính đủ giá bán theo cấu hình."],
        ["P1", "Kết nối approval matrix/request và hủy duyệt khi đổi giá", "Luồng kiểm soát thương mại hoàn chỉnh."],
        ["P2", "Waste rule mới, complexity, recovery, UoM/currency conversion", "Độ chính xác và bao phủ nghiệp vụ cao hơn."],
        ["P2", "Báo cáo breakdown, revision báo giá, audit và test tự động", "Vận hành, truy vết và bảo trì tốt."],
    ], [0.7, 3.05, 2.75], 8.8)

    doc.add_heading("15. Danh mục file code đã đối chiếu", level=1)
    doc.add_paragraph("Các kết luận trong tài liệu được rút ra từ code tại thời điểm rà soát. Danh mục dưới đây giúp đội phát triển tìm nhanh điểm cần sửa:")
    for t in [
        "dlm-erp/dl_sale/models/dl_quotation.py — model báo giá hiện tại.",
        "dlm-erp/dl_technical/models/dl_quotation_request.py — trạng thái RFQ, phân loại trading/manufactured và điều kiện resolve.",
        "dlm-erp/dl_technical/wizard/rfq_resolve_wizard.py — xử lý Product/BOM cho dòng RFQ.",
        "dlm-erp/dl_technical/models/dl_bom.py, dl_bom_line.py, dl_bom_line_mixin.py — định mức, hao hụt, snapshot giá và tổng chi phí vật tư.",
        "dlm-erp/dl_product/models/product_supplierinfo.py — phê duyệt và chọn bảng giá NCC đang áp dụng.",
        "dlm-erp/dl_config/models/pricing_rule.py — hiệu lực, revision và bảo vệ snapshot của rule.",
        "dlm-erp/dl_config/models/pricing_waste.py, pricing_operation.py, pricing_cost.py — hao hụt/thu hồi, công đoạn, chi phí chung.",
        "dlm-erp/dl_config/models/pricing_commercial.py — markup, giá sàn và chiết khấu theo nhóm khách.",
        "dlm-erp/dl_config/models/pricing_matrix.py, pricing_approval.py — đánh giá và xử lý phê duyệt.",
        "dlm-erp/dl_config/models/pricing_config.py — cấu hình tổng quan, VAT, làm tròn và ma trận/SLA cũ.",
    ]: bullet(doc, t)
    add_callout(doc, "Trạng thái quyết định:", "Tám quyết định kiến trúc/nghiệp vụ đã được chốt tại Decision log bên dưới. Các quyết định này thay thế những điểm còn để mở ở phiên bản 1.0 và là đầu vào chính thức cho kế hoạch P0/P1.", PALE_GREEN)

    doc.add_heading("16. Decision log - các quyết định đã chốt", level=1)
    doc.add_paragraph("Decision log có hiệu lực từ phiên bản 1.1 của tài liệu. Khi thay đổi một quyết định, phải tạo revision mới, ghi lý do và đánh giá ảnh hưởng tới báo giá/snapshot đã tồn tại.")
    decisions = [
        ["A1", "Nguồn hao hụt", "V3 dl.pricing.waste.rule là nguồn chân lý đích. P0/P1 tạm dùng waste_rate trên BOM; P2 chuyển resolver sang V3 và ngừng mở rộng dl.pricing.waste cũ.", "P0 kiến trúc / P2 chuyển đổi"],
        ["A2", "Snapshot giá", "Khi tạo báo giá, copy giá trị vô hướng vào dl.quotation.price.component; không compute lại theo BOM, sellerinfo hoặc rule sống.", "P0 - blocker"],
        ["B3", "per_batch/setup_fee", "Tính ở cấp dòng gia công; mỗi dòng là một lô. Phân bổ setup_fee về đơn giá bằng setup_fee/qty.", "P1"],
        ["B4", "CK/VAT hàng thương mại", "Giá gốc giữ nguyên list_price; vẫn tham gia chiết khấu và VAT cấp báo giá. Không áp BOM/chi phí gia công.", "P0"],
        ["B5", "Phân bổ CK để so sàn", "Phân bổ pro-rata theo subtotal chịu chiết khấu; so net_unit_price từng dòng gia công với floor_price.", "P1"],
        ["B6", "Dòng infeasible", "Chặn tạo báo giá cho toàn RFQ; Sales phải xác nhận/làm rõ lại phạm vi.", "P0 validation"],
        ["C7", "Chống tạo trùng", "Dùng search trước cho UX và UNIQUE DB trên quotation_request_id để chống race condition.", "P0 - blocker"],
        ["C8", "UoM/currency", "P0 chặn cứng khi khác UoM/currency bằng QTE-007; P1 triển khai quy đổi cơ bản, không để sai số âm thầm.", "P0 guard / P1 conversion"],
    ]
    add_table(doc, ["ID", "Chủ đề", "Quyết định chính thức", "Mốc"], decisions, [0.5, 1.2, 3.95, 0.85], 8.2)

    doc.add_heading("16.1. Hệ quả kiến trúc", level=2)
    for t in [
        "dl.quotation và các price component là hồ sơ bất biến của lần tính giá; BOM/rule là nguồn đầu vào, không phải nguồn hiển thị động của báo giá đã tạo.",
        "Model hao hụt cũ chỉ là cầu nối tạm thời. Không bổ sung recovery/complexity hoặc logic mới vào dl.pricing.waste.",
        "Giá thương mại có cùng pipeline chiết khấu/VAT với giá gia công, nhưng không đi qua pricing cost engine.",
        "Mọi phép so giá sàn dùng giá sau chiết khấu đã phân bổ ở cấp dòng, giúp nêu chính xác dòng nào tạo yêu cầu CEO.",
        "P0 ưu tiên tính đúng và chặn dữ liệu không tương thích; chưa quy đổi thì báo lỗi thay vì tự giả định.",
    ]:
        bullet(doc, t)

    doc.add_heading("16.2. Quy tắc ưu tiên khi có xung đột", level=2)
    add_table(doc, ["Xung đột", "Quy tắc xử lý"], [
        ["Tài liệu cũ khác Decision log", "Decision log phiên bản mới hơn được ưu tiên."],
        ["Rule đang active đổi sau khi tạo báo giá", "Báo giá cũ giữ snapshot; báo giá mới dùng revision mới."],
        ["list_price thay đổi sau khi tạo", "Dòng thương mại cũ giữ base_price snapshot."],
        ["RFQ/BOM thay đổi sau khi có báo giá", "Không sửa ngầm báo giá; tạo revision/tính lại và hủy duyệt cũ."],
        ["Có nhiều điều kiện duyệt", "Chọn cấp cao nhất; giữ toàn bộ reasons để audit."],
    ], [2.3, 4.2], 9.0)

    doc.add_heading("17. Kế hoạch triển khai P0 chi tiết", level=1)
    add_callout(doc, "Mục tiêu P0:", "Từ một RFQ confirmed, tạo đúng một báo giá draft có các dòng thương mại/gia công, snapshot giá nền bất biến, tổng tiền/chiết khấu/VAT cơ bản và validation an toàn. P0 chưa tính operation/cost adjustment V3 đầy đủ, nhưng thiết kế dữ liệu không được cản P1.", PALE_BLUE)

    doc.add_heading("17.1. Phạm vi P0", level=2)
    add_table(doc, ["Trong P0", "Chưa thuộc P0"], [
        ["Liên kết RFQ - quotation; chống trùng bằng DB", "6 phương pháp công đoạn và setup_fee thực tế"],
        ["Trading = list_price snapshot", "6 phương pháp cost adjustment"],
        ["Manufactured = BOM material cost snapshot/đơn vị", "Waste V3: complexity, recovery, scrap"],
        ["Chiết khấu/VAT header cơ bản", "Quy đổi UoM/currency đầy đủ (P1)"],
        ["Chặn infeasible, thiếu BOM/giá, UoM/currency khác", "Phê duyệt chi tiết theo line floor (P1 nếu cần breakdown đầy đủ)"],
        ["Price component snapshot cho nguồn giá P0", "Revision báo giá và báo cáo breakdown nâng cao"],
    ], [3.25, 3.25], 8.8)

    doc.add_heading("17.2. Thay đổi model", level=2)
    add_table(doc, ["Model", "Field/constraint P0", "Quy tắc"], [
        ["dl.quotation", "quotation_request_id (Many2one, required, index)", "Nguồn RFQ; ondelete=restrict."],
        ["dl.quotation", "SQL UNIQUE(quotation_request_id)", "Bảo đảm một RFQ chỉ có một báo giá P0."],
        ["dl.quotation", "company_id, pricing_date, currency_id", "Snapshot ngữ cảnh tính giá."],
        ["dl.quotation", "discount_pct, vat_pct", "Tỷ lệ header đã dùng."],
        ["dl.quotation", "amount_untaxed, discount_amount, amount_before_vat, vat_amount, amount_total", "Compute từ line snapshot, không đọc rule sống."],
        ["dl.quotation.line", "rfq_line_id, product_id, bom_id, line_type", "Truy vết và phân loại."],
        ["dl.quotation.line", "base_price, price_unit, floor_price, total_cost", "P0: base/price snapshot; floor có thể để 0 nếu profit rule chưa đưa vào P0."],
        ["dl.quotation.price.component", "quotation_line_id, component_type, source_model, source_id, source_revision", "Metadata truy vết nguồn."],
        ["dl.quotation.price.component", "qty, unit_price, rate, amount, no_discount", "Giá trị vô hướng bất biến."],
        ["dl.quotation.request", "quotation_id (computed hoặc related/search)", "Mở nhanh báo giá đã tạo; tránh lưu hai chiều nếu không cần."],
    ], [1.35, 2.55, 2.6], 8.2)

    doc.add_heading("17.3. Các loại price component P0", level=2)
    add_table(doc, ["component_type", "Áp dụng", "Giá trị snapshot"], [
        ["trading_base", "Dòng thương mại", "qty RFQ, list_price tại thời điểm tạo, subtotal."],
        ["material", "Mỗi dòng vật tư thô", "material_id/source sellerinfo, effective_qty quy về 1 output, unit_price, amount."],
        ["processed_material", "BOM con", "child_bom_id/version, child total snapshot, child product_qty, unit cost."],
        ["discount", "Header hoặc phân bổ line", "rate và amount thực dùng."],
        ["vat", "Header", "rate và amount thực dùng."],
    ], [1.45, 2.0, 3.05], 8.8)
    add_callout(doc, "Không dùng compute động:", "Sau khi component được tạo, thay đổi supplierinfo.price, is_applied, BOM.total_material_cost, list_price hoặc rule active không được làm thay đổi amount của báo giá. Tính lại phải là một action có chủ đích và tạo revision/hủy duyệt cũ ở phase sau.", PALE_YELLOW)

    doc.add_heading("17.4. Pricing service P0", level=2)
    doc.add_paragraph("Nên tách logic khỏi model action thành service/model abstract, ví dụ dl.quotation.pricing.service, để kiểm thử độc lập và tái sử dụng khi tính lại báo giá.")
    add_table(doc, ["Hàm gợi ý", "Trách nhiệm"], [
        ["build_context(rfq)", "Chốt company, currency, pricing_date, discount_pct, vat_pct và rule IDs/revisions P0."],
        ["validate_rfq(rfq, context)", "Kiểm tra status, infeasible, product/BOM, giá, UoM/currency và báo giá trùng."],
        ["price_trading(line, context)", "Snapshot list_price; trả line values + trading_base component."],
        ["price_bom(line, context, visited)", "Tính material cost đệ quy; phát hiện loop; chia child_bom.product_qty."],
        ["aggregate(quotation)", "Tính subtotal, discount, trước VAT, VAT và total từ snapshot."],
        ["create_from_rfq(rfq)", "Điều phối transaction và trả quotation/action."],
    ], [2.25, 4.25], 8.8)

    doc.add_heading("17.5. Trình tự action_create_quotation", level=2)
    p0_steps = [
        "ensure_one và kiểm quyền Sales/Admin; kiểm tra RFQ.status == confirmed.",
        "Search báo giá theo quotation_request_id để trả lỗi thân thiện trước khi insert.",
        "Build pricing context và chạy validate_rfq; chặn toàn RFQ nếu có bất kỳ dòng is_infeasible.",
        "Tạo dl.quotation header trong trạng thái draft.",
        "Tạo từng quotation line và component snapshot; trading không gọi BOM engine.",
        "Với manufactured, tính BOM theo đơn vị đầu ra; vật tư thiếu giá hoặc BOM con lỗi làm rollback toàn bộ.",
        "Tính discount/VAT. Hàng thương mại tham gia cùng discountable_amount và VAT.",
        "Ghi quotation_id/trạng thái quoted cho RFQ chỉ sau khi tất cả line/tổng tiền thành công.",
        "Bắt IntegrityError của UNIQUE để chuyển thành UserError 'RFQ đã có báo giá' trong race condition.",
        "Trả action mở form báo giá mới; không commit thủ công trong method Odoo.",
    ]
    for s in p0_steps:
        number(doc, s)

    doc.add_heading("17.6. Validation P0 theo Decision log", level=2)
    add_table(doc, ["Check", "Hành vi P0"], [
        ["RFQ có is_infeasible", "Chặn toàn bộ bằng lỗi nghiệp vụ; không bỏ qua dòng."],
        ["Trading thiếu/zero list_price", "Chặn QTE-003 tương đương nguồn giá bán không hợp lệ."],
        ["Manufactured thiếu BOM confirmed/locked", "Chặn QTE-002."],
        ["Supplierinfo không approved/is_applied hoặc price ≤ 0", "Chặn QTE-003."],
        ["BOM child product_qty ≤ 0 hoặc vòng lặp", "Chặn QTE-004."],
        ["UoM khác mà chưa có conversion P0", "Chặn QTE-007; không nhân trực tiếp."],
        ["Currency supplierinfo khác quotation/company", "Chặn QTE-007 trong P0."],
        ["quotation_request_id đã tồn tại", "Chặn bằng search và UNIQUE DB."],
    ], [2.7, 3.8], 8.8)

    doc.add_heading("17.7. File dự kiến thay đổi", level=2)
    add_table(doc, ["File", "Thay đổi P0"], [
        ["dl_sale/models/dl_quotation.py", "Mở rộng header/line, totals, constraint và action hỗ trợ."],
        ["dl_sale/models/quotation_pricing_service.py", "File mới: validate và tính snapshot trading/BOM."],
        ["dl_sale/models/quotation_price_component.py", "File mới: model component snapshot."],
        ["dl_sale/models/__init__.py", "Import service/component."],
        ["dl_sale/views/quotation_views.xml", "Hiển thị RFQ nguồn, totals, breakdown nội bộ và trạng thái."],
        ["dl_sale/security/ir.model.access.csv", "Quyền model component; che giá thành theo role."],
        ["dl_technical/models/dl_quotation_request.py", "Thêm action_create_quotation hoặc bridge gọi model dl_sale; xem lưu ý dependency."],
        ["dl_technical/views/quotation_request_views.xml", "Nút Tạo báo giá chỉ hiện khi confirmed."],
        ["dl_sale/__manifest__.py / dl_technical/__manifest__.py", "Chốt chiều dependency để tránh vòng phụ thuộc."],
    ], [2.65, 3.85], 8.5)
    add_callout(doc, "Lưu ý dependency Odoo:", "Hiện RFQ nằm trong dl_technical còn quotation nằm trong dl_sale. Không để dl_technical phụ thuộc dl_sale đồng thời dl_sale lại phụ thuộc dl_technical. Phương án khuyến nghị: dl_sale depends dl_technical và kế thừa dl.quotation.request để thêm action_create_quotation/nút view trong dl_sale; model gốc RFQ không import ngược quotation.", PALE_YELLOW)

    doc.add_heading("17.8. Definition of Done P0", level=2)
    for t in [
        "Tạo được báo giá từ RFQ mixed trading/manufactured bằng một thao tác.",
        "Mỗi RFQ chỉ tạo được một báo giá kể cả hai request đồng thời.",
        "Giá trading và material components giữ nguyên sau khi đổi list_price/supplierinfo/BOM nguồn.",
        "BOM con có product_qty > 1 được quy đổi unit cost đúng.",
        "Mọi trường hợp infeasible, thiếu giá/BOM, loop hoặc khác UoM/currency đều rollback và có thông báo rõ.",
        "Tổng trước chiết khấu, chiết khấu, trước VAT, VAT và tổng thanh toán đúng; trading có tham gia CK/VAT.",
        "RFQ chỉ chuyển quoted sau khi quotation, lines và components được tạo đầy đủ.",
        "Có automated tests tối thiểu cho TC01, TC02, TC03, TC09, TC11 và TC12.",
    ]:
        bullet(doc, t)

    doc.add_heading("17.9. Thứ tự thực hiện P0", level=2)
    add_table(doc, ["Thứ tự", "Work package", "Phụ thuộc"], [
        ["1", "Chốt chiều module dependency và schema quotation/component", "A1, A2, C7 đã chốt"],
        ["2", "Migration/model fields + SQL constraint + ACL", "WP1"],
        ["3", "Pricing service: validation, trading, BOM recursion", "WP2"],
        ["4", "action_create_quotation và transaction flow", "WP3"],
        ["5", "Views/nút thao tác/breakdown nội bộ", "WP4"],
        ["6", "Unit/integration/concurrency tests", "WP3-WP5"],
        ["7", "UAT với một RFQ hỗn hợp và dữ liệu giá thật", "WP6"],
    ], [0.7, 3.6, 2.2], 8.8)

    doc.core_properties.title = "Đặc tả tạo báo giá từ RFQ và BOM - Decision log và kế hoạch P0"
    doc.core_properties.subject = "DLM ERP - RFQ, BOM, giá vật tư và cấu hình tính giá"
    doc.core_properties.author = "DLM ERP Project"
    doc.core_properties.keywords = "RFQ, BOM, quotation, pricing, Odoo, decision log, P0"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
